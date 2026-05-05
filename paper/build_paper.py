from copy import deepcopy
from io import BytesIO
from pathlib import Path
import csv
import re
import zipfile

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = Path(r"C:\Users\ADMIN\Downloads\IRJMETSTemplate.docx")
OUTPUT_PATH = ROOT / "paper" / "HeartGaurd_IRJMETS_Paper.docx"
ASSETS_DIR = ROOT / "paper" / "assets"
FIGURE_ONE_PATH = ASSETS_DIR / "heartguard_architecture.png"
FIGURE_TWO_PATH = ASSETS_DIR / "heartguard_workflow.png"
HEADER_LOGO_INTERNAL_PATH = "word/media/image3.png"

MASTER_DATASET = ROOT / "data" / "final" / "heartguard_master_dataset_v1.csv"
CARDIO_DATASET = ROOT / "data" / "processed" / "cardiovascular_clean_v1.csv"
DIABETES_DATASET = ROOT / "data" / "processed" / "diabetes_binary_clean_v1.csv"
FRAMINGHAM_DATASET = ROOT / "data" / "processed" / "framingham_clean_v1.csv"
HEART_FAILURE_DATASET = ROOT / "data" / "processed" / "heart_failure_clean_v1.csv"
TRAINING_NOTEBOOK = ROOT / "notebooks" / "master_model_training.ipynb"


def load_font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\cambria.ttf" if not bold else r"C:\Windows\Fonts\cambriab.ttf",
        r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def set_run_font(run, name: str, size: int, bold: bool = False, italic: bool = False, color: RGBColor | None = None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    section_properties = deepcopy(body.sectPr) if body.sectPr is not None else None
    for child in list(body):
        body.remove(child)
    if section_properties is not None:
        body.append(section_properties)


def clear_story_part(part) -> None:
    element = part._element
    for child in list(element):
        element.remove(child)
    paragraph = part.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def simplify_template_margins(doc: Document) -> None:
    for section in doc.sections:
        for attribute in ("header", "first_page_header", "even_page_header", "footer", "first_page_footer", "even_page_footer"):
            if hasattr(section, attribute):
                clear_story_part(getattr(section, attribute))


def style_paragraph(
    paragraph,
    text: str,
    *,
    font_name: str = "Times New Roman",
    font_size: int = 10,
    bold: bool = False,
    italic: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    space_before: int = 0,
    space_after: int = 4,
    line_spacing: float = 1.05,
    color: RGBColor | None = None,
):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, font_name, font_size, bold=bold, italic=italic, color=color)
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing
    return paragraph


def add_paragraph(doc: Document, text: str, **kwargs):
    return style_paragraph(doc.add_paragraph(), text, **kwargs)


def add_title_block(doc: Document) -> None:
    add_paragraph(
        doc,
        "HEARTGUARD: PREVENTIVE INSIGHTS FOR HEART FAILURE READMISSION FORECASTING",
        font_name="Cambria",
        font_size=13,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=16,
        space_after=8,
        line_spacing=1.0,
    )
    add_paragraph(
        doc,
        "Author Name*1, Co-Author Name*2, Guide Name*3",
        font_name="Cambria",
        font_size=11,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
        line_spacing=1.0,
    )
    add_paragraph(
        doc,
        "*1Department of Computer Science, [Institute Name], [City], India",
        font_name="Cambria",
        font_size=10,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0,
        line_spacing=1.0,
    )
    add_paragraph(
        doc,
        "*2Department of Computer Science, [Institute Name], [City], India",
        font_name="Cambria",
        font_size=10,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0,
        line_spacing=1.0,
    )
    add_paragraph(
        doc,
        "*3Faculty Mentor, [Institute Name], [City], India",
        font_name="Cambria",
        font_size=10,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
        line_spacing=1.0,
    )
    add_paragraph(
        doc,
        "Author names and affiliations can be finalized before submission.",
        font_name="Times New Roman",
        font_size=9,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
        space_after=8,
        line_spacing=1.0,
    )


def add_section_heading(doc: Document, text: str) -> None:
    add_paragraph(
        doc,
        text,
        font_name="Cambria",
        font_size=12,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=4,
        space_after=4,
        line_spacing=1.0,
    )


def add_subheading(doc: Document, text: str) -> None:
    add_paragraph(
        doc,
        text,
        font_name="Cambria",
        font_size=10,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=2,
        space_after=2,
        line_spacing=1.0,
    )


def add_caption(doc: Document, text: str) -> None:
    add_paragraph(
        doc,
        text,
        font_name="Times New Roman",
        font_size=9,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=2,
        space_after=4,
        line_spacing=1.0,
    )


def configure_table_cell(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            set_run_font(run, "Times New Roman", 9, bold=bold)


def set_column_widths(table, widths):
    for column_index, width in enumerate(widths):
        for cell in table.columns[column_index].cells:
            cell.width = width


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_column_widths(table, widths)

    for index, header in enumerate(headers):
        configure_table_cell(table.rows[0].cells[index], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.CENTER if index in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
            configure_table_cell(cells[index], value, align=align)

    add_paragraph(doc, "", font_size=4, space_after=2, line_spacing=1.0)


def draw_rounded_box(draw, box, fill, outline, radius=28, width=4):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_centered_text(draw, box, text, font, fill, spacing=8):
    left, top, right, bottom = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]
    x = left + ((right - left) - text_width) / 2
    y = top + ((bottom - top) - text_height) / 2
    draw.multiline_text((x, y), text, font=font, fill=fill, spacing=spacing, align="center")


def draw_arrow(draw, start, end, color, width=8):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        arrow = [(x2, y2), (x2 - 20, y2 - 12), (x2 - 20, y2 + 12)]
    else:
        arrow = [(x2, y2), (x2 + 20, y2 - 12), (x2 + 20, y2 + 12)]
    draw.polygon(arrow, fill=color)


def create_header_logo_bytes() -> bytes:
    logo = Image.new("RGBA", (860, 420), (255, 255, 255, 0))
    draw = ImageDraw.Draw(logo)
    accent = "#0f766e"
    secondary = "#164e63"
    icon_fill = "#dcfce7"
    title_font = load_font(62, bold=True)
    subtitle_font = load_font(24, bold=False)

    draw.rounded_rectangle((24, 52, 220, 248), radius=42, fill=icon_fill, outline=accent, width=8)
    draw.polygon([(80, 132), (120, 92), (160, 132), (160, 180), (120, 214), (80, 180)], fill=accent)
    draw.line((110, 116, 130, 116), fill="white", width=10)
    draw.line((120, 106, 120, 126), fill="white", width=10)
    draw.text((262, 78), "HeartGuard", font=title_font, fill=secondary)
    draw.text((266, 168), "Preventive cardiovascular intelligence", font=subtitle_font, fill=accent)

    output = BytesIO()
    logo.save(output, format="PNG")
    return output.getvalue()


def create_figure_one() -> None:
    image = Image.new("RGB", (1600, 930), "#f5fbff")
    draw = ImageDraw.Draw(image)
    title_font = load_font(40, bold=True)
    box_title_font = load_font(30, bold=True)
    body_font = load_font(21)
    accent = "#0f766e"
    dark = "#12354a"
    fill = "#ffffff"
    muted = "#5f6f80"

    draw.text((72, 42), "HeartGuard forecasting architecture", font=title_font, fill=dark)

    boxes = {
        "sources": (70, 170, 400, 420),
        "harmonize": (470, 170, 800, 420),
        "forecast": (870, 170, 1200, 420),
        "support": (1260, 170, 1530, 420),
        "followup": (310, 540, 1285, 830),
    }

    for box in boxes.values():
        draw_rounded_box(draw, box, fill, accent)

    draw_centered_text(draw, boxes["sources"], "Clinical sources", box_title_font, dark)
    draw_centered_text(draw, (95, 290, 375, 405), "HF, Framingham,\ncardio, diabetes,\nand master data", body_font, muted)

    draw_centered_text(draw, boxes["harmonize"], "Feature harmonization", box_title_font, dark)
    draw_centered_text(draw, (495, 290, 775, 405), "Age, BMI, BP,\nglucose, cholesterol,\nEF, and creatinine", body_font, muted)

    draw_centered_text(draw, boxes["forecast"], "Explainable forecasting", box_title_font, dark)
    draw_centered_text(draw, (895, 290, 1175, 405), "Weighted score,\nconfidence, and\nsparse-input tolerance", body_font, muted)

    draw_centered_text(draw, boxes["support"], "Persistent history", box_title_font, dark)
    draw_centered_text(draw, (1285, 290, 1505, 405), "Saved records for\ndashboard continuity", body_font, muted)

    draw_centered_text(draw, boxes["followup"], "Patient and clinician support layer", box_title_font, dark)
    draw_centered_text(draw, (360, 670, 1235, 790), "Prediction, dashboard, diet, medication,\nmonitoring, mood, chat, and wellness support", body_font, muted)

    draw_arrow(draw, (400, 295), (470, 295), accent)
    draw_arrow(draw, (800, 295), (870, 295), accent)
    draw_arrow(draw, (1200, 295), (1260, 295), accent)
    draw_arrow(draw, (1035, 420), (1035, 540), accent)

    image.save(FIGURE_ONE_PATH)


def create_figure_two() -> None:
    image = Image.new("RGB", (1600, 900), "#fffdf8")
    draw = ImageDraw.Draw(image)
    title_font = load_font(40, bold=True)
    step_font = load_font(28, bold=True)
    detail_font = load_font(20)
    accent = "#f59e0b"
    dark = "#5b3a00"
    fill = "#ffffff"
    muted = "#6b5c43"

    draw.text((70, 42), "Operational workflow for preventive follow-up", font=title_font, fill=dark)

    steps = [
        ((70, 220, 320, 470), "1. Capture inputs", "Patient enters any\navailable 2 or more\nhealth indicators"),
        ((380, 220, 630, 470), "2. Estimate risk", "Backend returns risk,\nconfidence, and\ninterpretation"),
        ((690, 220, 940, 470), "3. Save history", "Prediction is appended\nto persistent Excel\nhistory"),
        ((1000, 220, 1250, 470), "4. Review trends", "Dashboard restores\nlatest records\nafter refresh"),
        ((1310, 220, 1560, 470), "5. Act on insight", "Diet, medication,\nmood, and wellness\nsupport adapt care"),
    ]

    for box, heading, detail in steps:
        draw_rounded_box(draw, box, fill, accent)
        draw_centered_text(draw, box, heading, step_font, dark)
        draw_centered_text(draw, (box[0] + 18, box[1] + 110, box[2] - 18, box[3] - 20), detail, detail_font, muted)

    for start_x, end_x in [(320, 380), (630, 690), (940, 1000), (1250, 1310)]:
        draw_arrow(draw, (start_x, 345), (end_x, 345), accent)

    footer_box = (280, 585, 1330, 790)
    draw_rounded_box(draw, footer_box, "#fff3d6", accent)
    draw_centered_text(
        draw,
        footer_box,
        "Design note: the deployed workflow favors explainability and sparse-input resilience.\n"
        "That makes the tool usable for preventive screening and readmission-aware follow-up,\n"
        "even before complete laboratory data are available.",
        detail_font,
        dark,
    )

    image.save(FIGURE_TWO_PATH)


def count_csv_rows(path: Path) -> int:
    with path.open("r", encoding="utf-8", newline="") as handle:
        return max(sum(1 for _ in handle) - 1, 0)


def count_csv_columns(path: Path) -> int:
    with path.open("r", encoding="utf-8", newline="") as handle:
        reader = csv.reader(handle)
        return len(next(reader))


def read_training_metrics() -> tuple[str, str]:
    text = TRAINING_NOTEBOOK.read_text(encoding="utf-8")
    accuracy_match = re.search(r"Accuracy:\s*([0-9.]+)", text)
    auc_match = re.search(r"ROC-AUC\s*:\s*([0-9.]+)", text)
    accuracy = accuracy_match.group(1) if accuracy_match else "0.0000"
    auc = auc_match.group(1) if auc_match else "0.0000"
    return accuracy, auc


def add_picture_block(doc: Document, image_path: Path, caption: str, width_inches: float) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    add_caption(doc, caption)


def replace_template_header_logo(docx_path: Path, replacement_bytes: bytes) -> None:
    temp_path = docx_path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(temp_path, "w") as target:
        for item in source.infolist():
            data = replacement_bytes if item.filename == HEADER_LOGO_INTERNAL_PATH else source.read(item.filename)
            target.writestr(item, data)
    temp_path.replace(docx_path)


def build_paper():
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    create_figure_one()
    create_figure_two()

    dataset_counts = {
        "master_rows": count_csv_rows(MASTER_DATASET),
        "master_cols": count_csv_columns(MASTER_DATASET),
        "cardio_rows": count_csv_rows(CARDIO_DATASET),
        "cardio_cols": count_csv_columns(CARDIO_DATASET),
        "diabetes_rows": count_csv_rows(DIABETES_DATASET),
        "diabetes_cols": count_csv_columns(DIABETES_DATASET),
        "framingham_rows": count_csv_rows(FRAMINGHAM_DATASET),
        "framingham_cols": count_csv_columns(FRAMINGHAM_DATASET),
        "heart_failure_rows": count_csv_rows(HEART_FAILURE_DATASET),
        "heart_failure_cols": count_csv_columns(HEART_FAILURE_DATASET),
    }
    accuracy, auc = read_training_metrics()

    abstract = (
        "HeartGuard is presented as a preventive decision-support platform that links cardiovascular screening, "
        "heart-failure-oriented feature analysis, and longitudinal follow-up in a single web workflow. The project "
        "uses curated clinical datasets and a browser-based FastAPI architecture to study how early risk signals can "
        "be transformed into patient-facing recommendations that are easier to understand and revisit after discharge. "
        "A harmonized project dataset with "
        f"{dataset_counts['master_rows']:,} records was prepared from cardiovascular, diabetes, Framingham, and heart failure sources, "
        "while the deployed application exposes an explainable weighted scoring engine that remains functional even when only a subset "
        "of variables is available. The system estimates risk, reports confidence, persists prediction history in an Excel-backed store, "
        "and connects the result to dashboard tracking, diet guidance, medication support, mood logging, monitoring, and wellness modules. "
        "Internal training artifacts from the project's master-model notebook demonstrated feasible predictive performance, while functional "
        "validation of the deployed prototype confirmed partial-input prediction, persistent history recovery, and continuity from assessment "
        "to support actions. The work shows that preventive insights for readmission-aware care can be made more accessible through an "
        "interpretable interface, though broader clinical validation and secure multi-user deployment remain necessary for production use."
    )

    introduction_paragraphs = [
        "Heart failure readmission remains one of the most persistent burdens in cardiovascular care because deterioration often begins before a patient recognizes its significance. The challenge is not only to detect elevated clinical risk, but also to maintain continuity between screening, follow-up behavior, and day-to-day self-management. Many patients are discharged with incomplete information, uncertain symptom awareness, and limited access to repeated consultation, which makes early preventive support especially valuable.",
        "Published machine learning studies have shown that readmission and cardiovascular deterioration can be anticipated from structured risk factors, yet many systems remain locked inside dataset-centric experiments or hospital dashboards that are difficult for end users to interpret. At the same time, lightweight preventive tools often sacrifice clinical grounding, persistence, or explainability. This creates a practical gap between predictive research and usable, everyday health-support workflows.",
        "HeartGuard addresses that gap by combining predictive reasoning with a continuity-first web experience. The project focuses on three connected goals: estimating heart-risk patterns from incomplete but meaningful health inputs, preserving those predictions for later review, and linking the newest assessment to support modules that encourage healthier follow-up behavior. In this framing, readmission forecasting is treated not as a stand-alone score, but as part of an actionable preventive pathway that connects risk awareness to sustained support.",
    ]

    methodology_paragraphs = [
        "The project was designed as a web-based prototype with a FastAPI backend, modular frontend pages, and persistent workbook storage for prediction history. Instead of forcing a single rigid intake form, the workflow accepts any meaningful combination of available health indicators and transforms them into a risk estimate. This makes the system more suitable for real-world preventive use, where users may know age, blood pressure, glucose, or medication history, but not every advanced cardiac value at the same time.",
        f"The data foundation combines multiple curated sources: a harmonized master dataset with {dataset_counts['master_rows']:,} records and {dataset_counts['master_cols']} columns, a cardiovascular screening dataset with {dataset_counts['cardio_rows']:,} records, a cleaned diabetes dataset with {dataset_counts['diabetes_rows']:,} records, the Framingham coronary risk dataset with {dataset_counts['framingham_rows']:,} records, and a heart failure dataset with {dataset_counts['heart_failure_rows']:,} records. The heart failure source is especially important because it includes post-discharge sensitive variables such as ejection fraction and serum creatinine, while the Framingham and cardiovascular sets help stabilize more general preventive features such as age, cholesterol, blood pressure, body mass index, smoking, and diabetes status.",
        f"Offline model-development notebooks in the project indicate that supervised learning is viable for this problem space; the master-model training notebook records an accuracy of {accuracy} and ROC-AUC of {auc}. However, the deployed web application intentionally uses an explainable weighted scoring strategy rather than a black-box ensemble. This choice was made for two reasons. First, the interface must stay operational when users provide only sparse inputs. Second, patient-facing tools benefit from transparent factor contributions and easier debugging. The result is a forecasting-oriented support layer that prioritizes interpretability and usability over raw model complexity in the current release.",
        "Each submitted payload is sanitized to remove empty fields, evaluated only on supported variables, and mapped into a weighted factor space. The application then computes an aggregate risk percentage and a confidence estimate linked to feature coverage. Confidence rises when richer clinical information is present, but the workflow remains available with at least two inputs. Every successful prediction is appended to a timestamped workbook entry so that the dashboard can recover the latest records after refresh and expose trends over time.",
    ]

    implementation_paragraphs = [
        "The frontend is organized around modular patient-support pages that extend the value of prediction into everyday use. In addition to the assessment form, the current project includes dashboard review, diet planning, chatbot assistance, medication support, mood logging, breathing exercises, SOS support, quiz-based education, and monitoring pages. This broader structure is important because readmission risk is not determined by one variable or one moment; it is shaped by adherence, symptom awareness, and continuity of healthy behavior.",
        "On the backend, FastAPI exposes health, prediction, and history endpoints. The deployed risk engine scores supported fields such as age, sex, body mass index, cholesterol, systolic and diastolic blood pressure, glucose, physical activity, smoking, hypertension, diabetes, heart rate, ejection fraction, and serum creatinine. These factors were selected because they combine preventive screening value with clinically meaningful heart-failure sensitivity, especially in cases where fluid balance, ventricular performance, or metabolic instability may influence deterioration risk.",
        "Persistence is implemented with OpenPyXL and an Excel history workbook. While this is not sufficient for a production hospital environment, it supports transparent prototype testing and demonstrates a useful design property: the dashboard can rehydrate prior predictions instead of resetting to an empty state after refresh. From a user-experience perspective, this continuity makes the platform feel substantially more reliable. It also supports doctor-patient review scenarios where the latest estimate, its confidence, and the number of inputs used should remain visible.",
        "To preserve explainability, the prediction view exposes risk level, confidence, input coverage, and a readable next-step message. The dashboard summarizes total predictions, average risk, high-risk cases, streak behavior, and the latest care note. Downstream support modules then personalize follow-up by adapting guidance to the newest saved profile. In practice, this means the forecasting output becomes a reusable context object for lifestyle and adherence recommendations rather than an isolated one-time score.",
    ]

    results_paragraphs = [
        "Functional validation of the deployed prototype confirmed that HeartGuard returns a usable result when at least two supported values are supplied. This directly improves accessibility for users who cannot complete a full cardiac intake form in one session. The same validation also showed that saved records are successfully appended to the workbook history and restored on the dashboard, which preserves the newest assessment after refresh and supports longitudinal monitoring rather than single-use estimation.",
        "The architecture also demonstrates a practical balance between machine-learning ambition and implementation realism. Internal notebook experiments suggest that supervised models can achieve meaningful discrimination on the harmonized dataset, yet the live application gains robustness by favoring sparse-input tolerance and transparent factor scoring. For preventive readmission-oriented support, this tradeoff is often preferable to a more opaque model that fails whenever one laboratory value is missing.",
        "Another important result is the continuity of action after prediction. Once a risk estimate is generated, the system does not stop at classification. The dashboard contextualizes the latest score, the diet module adjusts meal focus and avoid lists, and the surrounding support pages encourage medication adherence, breathing practice, and symptom awareness. This integrated flow strengthens the project's central claim that preventive insights are more useful when forecasting is linked to follow-up behavior.",
        "At the same time, the prototype reveals limits that should be acknowledged. Excel-backed persistence is acceptable for demonstration and single-team testing, but it lacks secure multi-user concurrency, access control, and clinical auditability. The deployed score is also readmission-aware rather than hospital-validated as a direct readmission probability. These constraints do not reduce the project's instructional value, but they define the boundary between a strong academic prototype and a production-grade healthcare platform.",
    ]

    testing_paragraphs = [
        "A separate testing perspective was applied to the deployed prototype because functional success in a healthcare-oriented workflow is not limited to raw prediction output. The application was checked across endpoint availability, minimum-input validation, persistence behavior, dashboard restoration, and continuity into support modules. These checks reflect the project's practical promise: even if the current forecasting layer is intentionally simplified, the user journey must remain stable and interpretable.",
        "API-level validation centered on the /api/health, /predict, and /predictions routes. The health endpoint verified service startup. The prediction endpoint was exercised with both incomplete and richer payloads to confirm correct input filtering, risk computation, confidence generation, and workbook appending. Invalid one-field submissions were expected to fail gracefully, while successful submissions were expected to produce both a visible assessment and a durable backend record. This behavior is central to readmission-aware preventive support because trust depends on consistent handling of partial, messy, real-world input.",
        "Client-side testing then examined the continuity path that begins after a prediction is received. The prediction page must expose the risk score, confidence, source label, and coverage count. The dashboard must reload saved history after refresh and turn that history into summary statistics such as average risk, high-risk counts, streak indicators, and latest care notes. Finally, downstream modules such as diet support must be able to reuse the latest saved profile rather than operating as disconnected pages. Together, these checks confirm that HeartGuard behaves like a workflow and not merely like a calculator.",
        "One useful safeguard already present in the frontend is the local fallback path when a non-client backend error occurs. Although backend storage is the preferred mode, a temporary local estimate still allows the interface to remain educational and responsive during development interruptions. From a testing perspective, this resilience matters because it preserves the continuity of interaction. From a safety perspective, the system still communicates that the source is local rather than backend-saved, which reduces the risk of presenting a fallback estimate as though it were a formally persisted clinical record.",
    ]

    future_scope = (
        "Future work should move in three directions. First, the project should reconnect its deployed interface to a validated model-serving pipeline so that training artifacts, calibration analysis, and prediction-time explanations are aligned end to end. Second, persistence should migrate from Excel storage to a secure database with role-based dashboards for patients, doctors, and administrators. Third, clinical evaluation should expand beyond functional verification to include curated heart-failure follow-up cohorts, threshold calibration, fairness analysis, and outcome-oriented testing that more directly measures readmission forecasting value."
    )

    conclusion = (
        "HeartGuard demonstrates that a preventive heart-support application can go beyond isolated risk scoring by combining explainable forecasting, persistent history, and actionable follow-up modules in one workflow. By integrating multi-source cardiovascular data, accepting incomplete inputs, preserving prior predictions, and translating the latest estimate into practical support, the system provides a credible foundation for readmission-aware preventive care. Its current value lies in transparency, usability, and continuity; its next step is rigorous clinical validation and secure deployment."
    )

    acknowledgements = (
        "The authors acknowledge the project guidance, peer review, and open clinical datasets that supported feature selection, prototyping, and experimentation for the HeartGuard platform."
    )

    references = [
        "[1] J.-S. Rachoin, K. Hunter, J. Varallo, and E. Cerceo, \"Impact of time from discharge to readmission on outcomes: an observational study from the US National Readmission Database,\" BMJ Open, vol. 14, 2024, Art. no. e085466.",
        "[2] S. Alajmani and H. Elazhary, \"Hospital Readmission Prediction using Machine Learning Techniques: A Comparative Study,\" International Journal of Advanced Computer Science and Applications, vol. 10, no. 4, pp. 212-220, 2019.",
        "[3] R. O. Baris and C. E. Tabit, \"Heart Failure Readmission Prevention Strategies-A Comparative Review of Medications, Devices, and Other Interventions,\" Journal of Clinical Medicine, vol. 14, no. 16, 2025, Art. no. 5894.",
        "[4] H. Hajishah et al., \"Evaluation of machine learning methods for prediction of heart failure mortality and readmission: meta-analysis,\" BMC Cardiovascular Disorders, vol. 25, 2025, Art. no. 264.",
        "[5] R. B. D'Agostino Sr. et al., \"General cardiovascular risk profile for use in primary care: the Framingham Heart Study,\" Circulation, vol. 117, no. 6, pp. 743-753, 2008.",
    ]

    annexure_paragraphs = [
        "The deployed HeartGuard prototype intentionally centers on variables that are understandable to patients, meaningful to clinicians, and resilient under incomplete reporting. Some variables, such as age, cholesterol, and blood pressure, support general cardiovascular screening. Others, such as ejection fraction and serum creatinine, provide stronger sensitivity to post-discharge heart-failure deterioration and therefore strengthen the platform's readmission-aware framing.",
        "This annexure summarizes the main variables currently scored by the backend and clarifies how they contribute to preventive interpretation. The list is useful for future teams because it links data collection, model reasoning, and user-facing explanation in one place.",
    ]

    doc = Document(TEMPLATE_PATH)
    clear_document_body(doc)
    simplify_template_margins(doc)

    add_title_block(doc)
    add_section_heading(doc, "ABSTRACT")
    add_paragraph(doc, abstract)
    add_paragraph(
        doc,
        "Keywords: heart failure readmission, preventive forecasting, cardiovascular risk assessment, FastAPI, explainable scoring, patient dashboard, longitudinal follow-up.",
        font_name="Cambria",
        font_size=10,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_after=5,
        line_spacing=1.0,
    )

    add_section_heading(doc, "I. INTRODUCTION")
    for paragraph in introduction_paragraphs:
        add_paragraph(doc, paragraph)

    add_subheading(doc, "Research Objectives")
    add_paragraph(
        doc,
        "The project objectives are to create a readmission-aware preventive workflow, preserve interpretable prediction history, and connect each new estimate to support modules that improve continuity after screening. These objectives are intentionally both technical and clinical in spirit: the system must be easy to use, but it must also respect the kinds of variables and follow-up behaviors that matter in cardiovascular care."
    )

    add_section_heading(doc, "II. METHODOLOGY")
    for paragraph in methodology_paragraphs:
        add_paragraph(doc, paragraph)

    add_caption(doc, "Table 1. Dataset foundations used to shape the HeartGuard forecasting workflow")
    add_table(
        doc,
        headers=["Dataset", "Records", "Columns", "Primary role in the project"],
        rows=[
            ["HeartGuard master dataset", f"{dataset_counts['master_rows']:,}", str(dataset_counts["master_cols"]), "Unified training and feature harmonization across cardiovascular domains"],
            ["Cardiovascular dataset", f"{dataset_counts['cardio_rows']:,}", str(dataset_counts["cardio_cols"]), "General heart-risk screening patterns and preventive profile stabilization"],
            ["Diabetes dataset", f"{dataset_counts['diabetes_rows']:,}", str(dataset_counts["diabetes_cols"]), "Metabolic comorbidity signals connected to adverse cardiovascular outcomes"],
            ["Framingham dataset", f"{dataset_counts['framingham_rows']:,}", str(dataset_counts["framingham_cols"]), "Classic long-term cardiovascular risk variables for primary-care context"],
            ["Heart failure dataset", f"{dataset_counts['heart_failure_rows']:,}", str(dataset_counts["heart_failure_cols"]), "Advanced cardiac indicators such as ejection fraction and serum creatinine"],
        ],
        widths=[Inches(1.65), Inches(0.8), Inches(0.8), Inches(3.15)],
    )

    add_picture_block(
        doc,
        FIGURE_ONE_PATH,
        "Figure 1: HeartGuard architecture from multi-source clinical features to persistent preventive support.",
        6.2,
    )

    add_section_heading(doc, "III. SYSTEM DESIGN AND IMPLEMENTATION")
    for paragraph in implementation_paragraphs:
        add_paragraph(doc, paragraph)

    add_caption(doc, "Table 2. Functional design mapping between forecasting output and support continuity")
    add_table(
        doc,
        headers=["Module", "Core input or state", "Output", "Value for follow-up care"],
        rows=[
            ["Prediction page", "Partial or complete patient metrics", "Risk score, confidence, risk label", "Creates the first preventive insight even with sparse data"],
            ["History workbook", "Saved backend record with timestamp", "Persistent longitudinal trace", "Prevents loss of context across refresh or later review"],
            ["Dashboard", "Latest saved prediction history", "Summary cards, trends, latest care note", "Makes recent status easy to revisit and compare"],
            ["Diet and wellness modules", "Newest saved profile and risk tier", "Personalized recommendations", "Turns risk awareness into daily action"],
        ],
        widths=[Inches(1.35), Inches(1.8), Inches(1.55), Inches(1.7)],
    )

    add_picture_block(
        doc,
        FIGURE_TWO_PATH,
        "Figure 2: Operational workflow from sparse health inputs to saved preventive follow-up guidance.",
        6.35,
    )

    add_section_heading(doc, "IV. RESULTS AND DISCUSSION")
    for paragraph in results_paragraphs:
        add_paragraph(doc, paragraph)

    add_caption(doc, "Table 3. Summary of observed prototype behavior during functional validation")
    add_table(
        doc,
        headers=["Validation focus", "Observed behavior", "Interpretation", "Prototype implication"],
        rows=[
            ["Sparse-input prediction", "Usable risk estimate produced with at least two values", "The interface remains accessible under incomplete data entry", "Supports early preventive screening instead of forcing full forms"],
            ["History persistence", "Saved records reloaded after dashboard refresh", "Context is preserved across sessions", "Improves continuity and trust in the tool"],
            ["Interpretability", "Risk label, coverage, and confidence stay visible", "Users can understand how complete or limited a result is", "Makes patient-facing forecasting safer and easier to explain"],
            ["Action continuity", "Prediction state is reused by dashboard and lifestyle modules", "Risk insight informs downstream support behavior", "Connects forecasting to practical follow-up care"],
        ],
        widths=[Inches(1.35), Inches(1.85), Inches(1.55), Inches(1.65)],
    )

    add_section_heading(doc, "V. TESTING AND EVALUATION")
    for paragraph in testing_paragraphs:
        add_paragraph(doc, paragraph)

    add_caption(doc, "Table 4. Targeted testing observations for the deployed HeartGuard prototype")
    add_table(
        doc,
        headers=["Scenario", "Expected outcome", "Observed result", "Why it matters"],
        rows=[
            ["Two-field prediction", "System should allow forecasting with sparse but valid input", "Prediction returned with risk, confidence, and saved record", "Supports early screening when complete data are unavailable"],
            ["Single-field submission", "System should reject insufficient evidence", "Graceful client-facing validation error", "Prevents misleading estimates from underspecified input"],
            ["Dashboard refresh", "Latest saved records should remain visible", "History restored from workbook-backed endpoint", "Maintains longitudinal continuity and user trust"],
            ["Backend interruption", "Frontend should degrade safely", "Local fallback estimate remains clearly labeled", "Preserves educational use without hiding data-source limitations"],
        ],
        widths=[Inches(1.45), Inches(1.8), Inches(1.55), Inches(1.7)],
    )

    add_section_heading(doc, "VI. LIMITATIONS AND FUTURE SCOPE")
    add_paragraph(
        doc,
        "The current HeartGuard release should be interpreted as an academic and functional prototype, not a hospital-certified readmission engine. Its strengths lie in explainability, continuity, and preventive usability, while its present limitations include workbook persistence, lack of authenticated multi-user workflows, absence of prospective clinical validation, and simplified deployment assumptions. These limits are typical at the prototype stage, but they should be addressed before the platform is used in a higher-stakes environment."
    )
    add_paragraph(doc, future_scope)

    add_section_heading(doc, "VII. CONCLUSION")
    add_paragraph(doc, conclusion)

    add_section_heading(doc, "ACKNOWLEDGEMENTS")
    add_paragraph(doc, acknowledgements)

    add_section_heading(doc, "REFERENCES")
    for reference in references:
        add_paragraph(doc, reference, font_size=9, line_spacing=1.0, space_after=2)

    doc.add_page_break()
    add_section_heading(doc, "ANNEXURE: CORE HEARTGUARD VARIABLES")
    for paragraph in annexure_paragraphs:
        add_paragraph(doc, paragraph)

    add_caption(doc, "Table 5. Key variables used by the deployed HeartGuard scoring workflow")
    add_table(
        doc,
        headers=["Variable", "Clinical meaning", "Prototype role", "Interpretation in follow-up support"],
        rows=[
            ["Age", "Baseline cardiovascular vulnerability", "General preventive signal", "Helps contextualize long-term risk and prioritization"],
            ["BMI", "Body-weight related cardiometabolic stress", "Lifestyle-sensitive input", "Supports diet and activity recommendations"],
            ["Systolic and diastolic blood pressure", "Hemodynamic burden and hypertension severity", "High-weight screening factors", "Guides caution around elevated cardiac strain"],
            ["Cholesterol", "Atherogenic burden", "Preventive screening input", "Contributes to lifestyle and long-term risk messaging"],
            ["Glucose and diabetes status", "Metabolic instability and vascular risk", "Comorbidity-sensitive input", "Connects prediction to diet and chronic-disease awareness"],
            ["Smoking and physical activity", "Behavioral exposure and protective habits", "Modifiable lifestyle indicators", "Used to personalize preventive guidance"],
            ["Heart rate", "Physiologic stress or instability", "Secondary dynamic marker", "Adds context to overall current-state interpretation"],
            ["Ejection fraction", "Cardiac pumping performance", "Heart-failure-sensitive variable", "Important for readmission-aware deterioration assessment"],
            ["Serum creatinine", "Renal function and fluid-management sensitivity", "Advanced heart-failure feature", "Useful for higher-risk post-discharge interpretation"],
        ],
        widths=[Inches(1.35), Inches(1.75), Inches(1.55), Inches(1.85)],
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved paper to {OUTPUT_PATH}")


if __name__ == "__main__":
    build_paper()
