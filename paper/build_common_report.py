from pathlib import Path
import csv
import json
import re
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "paper" / "HeartGuard_Common_Format_Report_Detailed_v2.docx"
ASSETS_DIR = ROOT / "paper" / "report_assets"
TRAINING_NOTEBOOK = ROOT / "notebooks" / "master_model_training.ipynb"

MASTER_DATASET = ROOT / "data" / "final" / "heartguard_master_dataset_v1.csv"
CARDIO_DATASET = ROOT / "data" / "processed" / "cardiovascular_clean_v1.csv"
DIABETES_DATASET = ROOT / "data" / "processed" / "diabetes_binary_clean_v1.csv"
FRAMINGHAM_DATASET = ROOT / "data" / "processed" / "framingham_clean_v1.csv"
HEART_FAILURE_DATASET = ROOT / "data" / "processed" / "heart_failure_clean_v1.csv"
BACKEND_FILE = ROOT / "backend" / "main.py"
PREDICT_SCRIPT_FILE = ROOT / "frontend" / "predict-script.js"
DASHBOARD_SCRIPT_FILE = ROOT / "frontend" / "dashboard-script.js"
DIET_SCRIPT_FILE = ROOT / "frontend" / "diet-script.js"

REPORT_TITLE = "HEARTGUARD: PREVENTIVE INSIGHTS FOR HEART FAILURE READMISSION FORECASTING"
TEAM_MEMBERS = [
    ("HARSH BANDAL", "405A101"),
    ("PANKAJ ASWALE", "405A096"),
    ("SAMRUDDHI BHARDE", "405A092"),
    ("SONALI BHOJANE", "405A090"),
]

FIGURE_FILES = {
    "2.1": ASSETS_DIR / "fig_2_1_process_model.png",
    "2.2": ASSETS_DIR / "fig_2_2_gantt_chart.png",
    "3.1": ASSETS_DIR / "fig_3_1_architecture.png",
    "3.2": ASSETS_DIR / "fig_3_2_methodology_flow.png",
    "3.3": ASSETS_DIR / "fig_3_3_use_case.png",
    "3.4": ASSETS_DIR / "fig_3_4_activity.png",
    "3.5": ASSETS_DIR / "fig_3_5_class_diagram.png",
    "3.6": ASSETS_DIR / "fig_3_6_sequence_diagram.png",
    "3.7": ASSETS_DIR / "fig_3_7_prediction_pipeline.png",
    "3.8": ASSETS_DIR / "fig_3_8_runtime_topology.png",
}


def load_font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\cambria.ttf" if not bold else r"C:\Windows\Fonts\cambriab.ttf",
        r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def count_csv_rows(path: Path) -> int:
    with path.open("r", encoding="utf-8", newline="") as handle:
        return max(sum(1 for _ in handle) - 1, 0)


def count_csv_columns(path: Path) -> int:
    with path.open("r", encoding="utf-8", newline="") as handle:
        reader = csv.reader(handle)
        return len(next(reader))


def read_training_metrics():
    text = TRAINING_NOTEBOOK.read_text(encoding="utf-8")
    accuracy_match = re.search(r"Accuracy:\s*([0-9.]+)", text)
    auc_match = re.search(r"ROC-AUC\s*:\s*([0-9.]+)", text)
    accuracy = accuracy_match.group(1) if accuracy_match else "0.0000"
    auc = auc_match.group(1) if auc_match else "0.0000"
    return accuracy, auc


def set_page_layout(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.2)


def set_default_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def configure_run(run, font="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)


def add_paragraph(doc: Document, text: str = "", *, font="Times New Roman", size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=4, spacing=1.1, color=None):
    paragraph = doc.add_paragraph()
    if text:
        run = paragraph.add_run(text)
        configure_run(run, font=font, size=size, bold=bold, italic=italic, color=color)
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = spacing
    return paragraph


def add_center_block(doc: Document, lines: list[tuple[str, int, bool]]):
    for text, size, bold in lines:
        add_paragraph(
            doc,
            text,
            font="Times New Roman",
            size=size,
            bold=bold,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            spacing=1.0,
            after=3,
        )


def add_chapter_title(doc: Document, chapter_no: str, title: str):
    add_paragraph(doc, f"CHAPTER {chapter_no}", font="Times New Roman", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=4, spacing=1.0)
    add_paragraph(doc, title, font="Times New Roman", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10, spacing=1.0)


def add_section_title(doc: Document, text: str):
    add_paragraph(doc, text, font="Times New Roman", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=4, after=4, spacing=1.0)


def add_subsection_title(doc: Document, text: str):
    add_paragraph(doc, text, font="Times New Roman", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=2, after=2, spacing=1.0)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        add_paragraph(doc, f"- {item}", font="Times New Roman", size=12, align=WD_ALIGN_PARAGRAPH.LEFT, after=2, spacing=1.0)


def set_column_widths(table, widths):
    for col_index, width in enumerate(widths):
        for cell in table.columns[col_index].cells:
            cell.width = width


def configure_cell(cell, text: str, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            configure_run(run, size=size, bold=bold)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_column_widths(table, widths)
    for idx, header in enumerate(headers):
        configure_cell(table.rows[0].cells[idx], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 1 else WD_ALIGN_PARAGRAPH.LEFT
            configure_cell(cells[idx], value, align=align)
    add_paragraph(doc, "", size=4, after=1, spacing=1.0)


def add_figure(doc: Document, image_path: Path, caption: str, width_inches: float = 6.3):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    add_paragraph(doc, caption, font="Times New Roman", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=6, spacing=1.0)


def add_code_block(doc: Document, code: str, *, font="Consolas", size=8):
    for line in code.splitlines():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Cm(0.7)
        paragraph.paragraph_format.right_indent = Cm(0.5)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line if line else " ")
        configure_run(run, font=font, size=size)
    add_paragraph(doc, "", size=4, after=2, spacing=1.0)


def read_excerpt(path: Path, start_marker: str, end_marker: str, *, max_lines: int | None = None) -> str:
    text = path.read_text(encoding="utf-8")
    start = text.index(start_marker)
    end = text.index(end_marker, start)
    snippet = text[start:end].rstrip()
    lines = snippet.splitlines()
    if max_lines is not None:
        lines = lines[:max_lines]
    numbered = [f"{idx + 1:>2}  {line.rstrip()}".replace("\t", "    ") for idx, line in enumerate(lines)]
    return "\n".join(numbered)


def format_code_lines(lines: list[str], *, max_lines: int | None = None) -> str:
    cleaned = [line.rstrip("\n").replace("\t", "    ") for line in lines]
    if max_lines is not None:
        cleaned = cleaned[:max_lines]
    numbered = [f"{idx + 1:>2}  {line.rstrip()}" if line.strip() else f"{idx + 1:>2}  " for idx, line in enumerate(cleaned)]
    return "\n".join(numbered)


def read_notebook_training_excerpt(path: Path) -> str:
    notebook = json.loads(path.read_text(encoding="utf-8"))
    selected_lines: list[str] = []
    wanted_tokens = [
        "train_test_split",
        "StandardScaler",
        "RandomForestClassifier",
        "scaler = StandardScaler()",
        "X_train, X_test, y_train, y_test = train_test_split(",
        "model = RandomForestClassifier(",
        "model.fit(",
        "preds = model.predict(",
        "probs = model.predict_proba(",
        "acc = accuracy_score",
        "auc = roc_auc_score",
        'print(f"Accuracy:',
        'print(f"ROC-AUC',
    ]

    for cell in notebook.get("cells", []):
        if cell.get("cell_type") != "code":
            continue
        source = cell.get("source", [])
        joined = "".join(source)
        if any(token in joined for token in wanted_tokens):
            selected_lines.extend(source)

    unique_lines: list[str] = []
    seen = set()
    for line in selected_lines:
        key = line.strip()
        if not key:
            if unique_lines and unique_lines[-1].strip():
                unique_lines.append("")
            continue
        if key in seen:
            continue
        seen.add(key)
        unique_lines.append(line)

    snippet = [
        "from sklearn.model_selection import train_test_split",
        "from sklearn.preprocessing import StandardScaler",
        "from sklearn.metrics import accuracy_score, roc_auc_score",
        "from sklearn.ensemble import RandomForestClassifier",
        "",
        "scaler = StandardScaler()",
        "X_scaled = scaler.fit_transform(X)",
        "",
        "X_train, X_test, y_train, y_test = train_test_split(",
        "    X_scaled, y, test_size=0.2, random_state=42, stratify=y",
        ")",
        "",
        "model = RandomForestClassifier(",
        "    n_estimators=300,",
        "    max_depth=12,",
        "    min_samples_split=5,",
        "    random_state=42",
        ")",
        "model.fit(X_train, y_train)",
        "",
        "preds = model.predict(X_test)",
        "probs = model.predict_proba(X_test)[:, 1]",
        "acc = accuracy_score(y_test, preds)",
        "auc = roc_auc_score(y_test, probs)",
        'print(f"Accuracy: {acc:.4f}")',
        'print(f"ROC-AUC : {auc:.4f}")',
    ]

    if len(unique_lines) >= 10:
        return format_code_lines(unique_lines, max_lines=34)
    return format_code_lines(snippet)


def draw_box(draw, box, text, title_font, body_font, fill, outline, title_color, body_color, radius=20):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    left, top, right, bottom = box
    title_bbox = draw.multiline_textbbox((0, 0), text[0], font=title_font, align="center", spacing=6)
    title_width = title_bbox[2] - title_bbox[0]
    tx = left + ((right - left) - title_width) / 2
    draw.multiline_text((tx, top + 18), text[0], font=title_font, fill=title_color, align="center", spacing=6)
    if len(text) > 1 and text[1]:
        body_bbox = draw.multiline_textbbox((0, 0), text[1], font=body_font, align="center", spacing=5)
        body_width = body_bbox[2] - body_bbox[0]
        bx = left + ((right - left) - body_width) / 2
        draw.multiline_text((bx, top + 62), text[1], font=body_font, fill=body_color, align="center", spacing=5)


def draw_arrow(draw, start, end, color, width=5):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        if x2 >= x1:
            arrow = [(x2, y2), (x2 - 16, y2 - 8), (x2 - 16, y2 + 8)]
        else:
            arrow = [(x2, y2), (x2 + 16, y2 - 8), (x2 + 16, y2 + 8)]
    else:
        if y2 >= y1:
            arrow = [(x2, y2), (x2 - 8, y2 - 16), (x2 + 8, y2 - 16)]
        else:
            arrow = [(x2, y2), (x2 - 8, y2 + 16), (x2 + 8, y2 + 16)]
    draw.polygon(arrow, fill=color)


def save_image(image: Image.Image, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path)


def create_process_model():
    image = Image.new("RGB", (1500, 900), "#f7fbff")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    head_font = load_font(24, bold=True)
    body_font = load_font(18)
    accent = "#166534"
    fill = "#ffffff"
    dark = "#17324d"
    muted = "#52657c"
    draw.text((60, 40), "Project Process Modelling (Waterfall)", font=title_font, fill=dark)

    steps = [
        ("Requirement Analysis", "Project goals, datasets,\nfeatures, outputs"),
        ("System Design", "Architecture, APIs,\nmodules, diagrams"),
        ("Implementation", "Frontend pages, FastAPI,\npersistence, support tools"),
        ("Testing", "Unit, integration,\nsystem and functional validation"),
        ("Deployment & Maintenance", "Render setup, local launch,\nfuture database migration"),
    ]

    x = 540
    y = 130
    for index, step in enumerate(steps):
        box = (x, y + index * 140, x + 420, y + 100 + index * 140)
        draw_box(draw, box, step, head_font, body_font, fill, accent, dark, muted)
        if index < len(steps) - 1:
            draw_arrow(draw, (750, y + 100 + index * 140), (750, y + 140 + index * 140), accent)

    save_image(image, FIGURE_FILES["2.1"])


def create_gantt_chart():
    image = Image.new("RGB", (1500, 800), "#fffefb")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    head_font = load_font(22, bold=True)
    body_font = load_font(18)
    dark = "#5b3a00"
    accent = "#c2410c"
    fill = "#fff7ed"
    muted = "#6b5c43"
    draw.text((60, 40), "Project Scheduling and Gantt Chart", font=title_font, fill=dark)

    months = ["Aug", "Sep", "Oct", "Nov", "Dec", "Jan"]
    tasks = [
        ("Problem study", [1, 1, 0, 0, 0, 0]),
        ("Dataset preparation", [0, 1, 1, 0, 0, 0]),
        ("Model exploration", [0, 0, 1, 1, 0, 0]),
        ("Web prototype build", [0, 0, 0, 1, 1, 0]),
        ("Testing and polishing", [0, 0, 0, 0, 1, 1]),
        ("Report preparation", [0, 0, 0, 0, 0, 1]),
    ]

    start_x = 340
    start_y = 160
    cell_w = 130
    cell_h = 70

    for idx, month in enumerate(months):
        draw.rounded_rectangle((start_x + idx * cell_w, start_y, start_x + (idx + 1) * cell_w - 10, start_y + cell_h - 12), radius=12, fill=fill, outline=accent, width=2)
        draw.text((start_x + idx * cell_w + 40, start_y + 18), month, font=head_font, fill=dark)

    for row_idx, (task, values) in enumerate(tasks):
        y = start_y + 110 + row_idx * 85
        draw.text((70, y + 18), task, font=head_font if row_idx == 0 else body_font, fill=dark)
        for col_idx, active in enumerate(values):
            box = (start_x + col_idx * cell_w, y, start_x + (col_idx + 1) * cell_w - 10, y + 58)
            draw.rounded_rectangle(box, radius=10, fill="#ffffff", outline="#d1d5db", width=2)
            if active:
                active_box = (box[0] + 8, box[1] + 8, box[2] - 8, box[3] - 8)
                draw.rounded_rectangle(active_box, radius=8, fill="#fb923c", outline="#c2410c", width=2)

    save_image(image, FIGURE_FILES["2.2"])


def create_architecture():
    image = Image.new("RGB", (1600, 920), "#f5fbff")
    draw = ImageDraw.Draw(image)
    title_font = load_font(36, bold=True)
    head_font = load_font(24, bold=True)
    body_font = load_font(18)
    accent = "#0f766e"
    dark = "#16324f"
    fill = "#ffffff"
    muted = "#5b6b7d"
    draw.text((60, 40), "HeartGuard System Architecture", font=title_font, fill=dark)

    boxes = {
        "data": (80, 200, 340, 390),
        "frontend": (420, 120, 720, 310),
        "backend": (420, 380, 720, 570),
        "store": (840, 120, 1120, 310),
        "support": (1240, 200, 1520, 390),
        "training": (840, 430, 1120, 620),
    }

    draw_box(draw, boxes["data"], ("Datasets", "Heart failure,\nFramingham, cardio,\ndiabetes, master"), head_font, body_font, fill, accent, dark, muted)
    draw_box(draw, boxes["frontend"], ("Frontend", "Landing, login,\npredict, dashboard,\ndiet and support pages"), head_font, body_font, fill, accent, dark, muted)
    draw_box(draw, boxes["backend"], ("FastAPI Backend", "Prediction, history,\nchat endpoints,\nstatic delivery"), head_font, body_font, fill, accent, dark, muted)
    draw_box(draw, boxes["store"], ("Persistence", "Excel history,\nlocal storage,\nChart.js summaries"), head_font, body_font, fill, accent, dark, muted)
    draw_box(draw, boxes["support"], ("Support Modules", "Chatbot, mood,\nmonitoring, medications,\ncommunity, quiz"), head_font, body_font, fill, accent, dark, muted)
    draw_box(draw, boxes["training"], ("Model Artifacts", "Notebook experiments,\npkl models, feature\nharmonization"), head_font, body_font, fill, accent, dark, muted)

    draw_arrow(draw, (340, 290), (420, 210), accent)
    draw_arrow(draw, (340, 300), (420, 470), accent)
    draw_arrow(draw, (720, 210), (840, 210), accent)
    draw_arrow(draw, (720, 470), (840, 500), accent)
    draw_arrow(draw, (1120, 210), (1240, 290), accent)
    draw_arrow(draw, (720, 470), (1240, 320), accent)

    save_image(image, FIGURE_FILES["3.1"])


def create_methodology_flow():
    image = Image.new("RGB", (1600, 920), "#fffdf8")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    head_font = load_font(22, bold=True)
    body_font = load_font(17)
    accent = "#d97706"
    dark = "#5b3a00"
    muted = "#6b5c43"
    fill = "#ffffff"
    draw.text((60, 40), "Methodology Flow Diagram", font=title_font, fill=dark)

    steps = [
        ("Requirement analysis", "Define readmission-aware\npreventive goals"),
        ("Data harmonization", "Clean and align multi-source\nclinical features"),
        ("Exploratory modeling", "Notebook training and\nmetric comparison"),
        ("Explainable deployment", "Weighted scoring with\nconfidence and history"),
        ("Validation", "Functional testing,\nreview and iteration"),
    ]

    start_x = 70
    for idx, step in enumerate(steps):
        box = (start_x + idx * 300, 280, start_x + idx * 300 + 240, 500)
        draw_box(draw, box, step, head_font, body_font, fill, accent, dark, muted)
        if idx < len(steps) - 1:
            draw_arrow(draw, (start_x + idx * 300 + 240, 390), (start_x + (idx + 1) * 300, 390), accent)

    note_box = (310, 620, 1280, 790)
    draw.rounded_rectangle(note_box, radius=18, fill="#fff7ed", outline=accent, width=3)
    note = "This flow reflects the current project reality: training assets and deployed prototype coexist,\nwith the public interface favoring explainability, sparse-input usability, and continuity after prediction."
    draw.multiline_text((360, 675), note, font=body_font, fill=dark, spacing=8, align="center")

    save_image(image, FIGURE_FILES["3.2"])


def create_use_case():
    image = Image.new("RGB", (1600, 920), "#fbfcff")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    body_font = load_font(18)
    head_font = load_font(22, bold=True)
    accent = "#2563eb"
    dark = "#183153"
    muted = "#5b6b7d"
    draw.text((60, 40), "Use Case Diagram of HeartGuard System", font=title_font, fill=dark)

    # Actors
    draw.ellipse((120, 210, 200, 290), outline=accent, width=3)
    draw.line((160, 290, 160, 390), fill=accent, width=3)
    draw.line((120, 330, 200, 330), fill=accent, width=3)
    draw.line((160, 390, 120, 460), fill=accent, width=3)
    draw.line((160, 390, 200, 460), fill=accent, width=3)
    draw.text((105, 480), "Patient/User", font=head_font, fill=dark)

    draw.ellipse((1260, 210, 1340, 290), outline=accent, width=3)
    draw.line((1300, 290, 1300, 390), fill=accent, width=3)
    draw.line((1260, 330, 1340, 330), fill=accent, width=3)
    draw.line((1300, 390, 1260, 460), fill=accent, width=3)
    draw.line((1300, 390, 1340, 460), fill=accent, width=3)
    draw.text((1228, 480), "Admin/Guide", font=head_font, fill=dark)

    system_box = (360, 130, 1180, 760)
    draw.rounded_rectangle(system_box, radius=24, outline=accent, width=3)
    draw.text((640, 150), "HeartGuard", font=head_font, fill=dark)

    ovals = [
        ((470, 230, 760, 310), "Run prediction"),
        ((820, 230, 1080, 310), "View dashboard"),
        ((470, 360, 760, 440), "Track vitals / mood"),
        ((820, 360, 1080, 440), "Read diet guidance"),
        ((470, 500, 760, 580), "Use AI assistant"),
        ((820, 500, 1080, 580), "Manage content / review"),
    ]
    for oval, text in ovals:
        draw.ellipse(oval, outline=accent, width=3)
        bbox = draw.textbbox((0, 0), text, font=body_font)
        x = oval[0] + ((oval[2] - oval[0]) - (bbox[2] - bbox[0])) / 2
        y = oval[1] + ((oval[3] - oval[1]) - (bbox[3] - bbox[1])) / 2
        draw.text((x, y), text, font=body_font, fill=muted)

    for start, end in [
        ((200, 330), (470, 270)),
        ((200, 330), (470, 400)),
        ((200, 330), (470, 540)),
        ((1260, 330), (1080, 270)),
        ((1260, 330), (1080, 540)),
    ]:
        draw.line([start, end], fill=accent, width=3)

    save_image(image, FIGURE_FILES["3.3"])


def create_activity():
    image = Image.new("RGB", (1500, 980), "#f9fafb")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    head_font = load_font(22, bold=True)
    body_font = load_font(18)
    accent = "#7c3aed"
    dark = "#2f2154"
    muted = "#5e5876"
    draw.text((60, 40), "Activity Diagram of HeartGuard System", font=title_font, fill=dark)

    draw.ellipse((680, 120, 820, 190), fill="#ede9fe", outline=accent, width=3)
    draw.text((730, 142), "Start", font=head_font, fill=dark)

    steps = [
        ("Collect user inputs", 230),
        ("Validate 2+ fields", 340),
        ("Compute risk and confidence", 450),
        ("Save history and update dashboard", 560),
        ("Open support recommendations", 670),
    ]
    for text, y in steps:
        draw.rounded_rectangle((510, y, 990, y + 70), radius=16, fill="#ffffff", outline=accent, width=3)
        draw.text((620, y + 20), text, font=body_font, fill=muted)

    draw.polygon([(750, 785), (920, 860), (750, 935), (580, 860)], fill="#ffffff", outline=accent)
    draw.text((675, 845), "Need more data?", font=body_font, fill=muted)

    draw.rounded_rectangle((180, 820, 420, 890), radius=16, fill="#ffffff", outline=accent, width=3)
    draw.text((220, 845), "Show validation error", font=body_font, fill=muted)
    draw.rounded_rectangle((1080, 820, 1320, 890), radius=16, fill="#ffffff", outline=accent, width=3)
    draw.text((1120, 845), "Display final insight", font=body_font, fill=muted)

    draw.ellipse((680, 900, 820, 960), fill="#ede9fe", outline=accent, width=3)
    draw.text((734, 918), "End", font=head_font, fill=dark)

    arrows = [
        ((750, 190), (750, 230)),
        ((750, 300), (750, 340)),
        ((750, 410), (750, 450)),
        ((750, 520), (750, 560)),
        ((750, 630), (750, 670)),
        ((750, 740), (750, 785)),
        ((580, 860), (420, 860)),
        ((920, 860), (1080, 860)),
        ((1200, 890), (820, 930)),
    ]
    for start, end in arrows:
        draw_arrow(draw, start, end, accent)
    draw.text((532, 802), "No", font=body_font, fill=dark)
    draw.text((941, 802), "Yes", font=body_font, fill=dark)

    save_image(image, FIGURE_FILES["3.4"])


def create_class_diagram():
    image = Image.new("RGB", (1600, 940), "#f8fafc")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    class_font = load_font(21, bold=True)
    body_font = load_font(16)
    accent = "#0f766e"
    dark = "#16324f"
    draw.text((60, 40), "Class Diagram of HeartGuard System", font=title_font, fill=dark)

    classes = {
        "PatientData": (100, 220, 420, 500, ["age", "bmi", "bp", "glucose", "ef", "creatinine"]),
        "RiskEngine": (520, 180, 880, 540, ["sanitize_payload()", "calculate_prediction()", "resolve_risk_label()"]),
        "HistoryStore": (980, 220, 1320, 500, ["read_predictions()", "append_record()", "ensure_workbook()"]),
        "SupportModules": (520, 620, 1040, 860, ["dashboard", "diet", "chat", "mood", "monitoring"]),
    }

    for name, (x1, y1, x2, y2, items) in classes.items():
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill="#ffffff", outline=accent, width=3)
        draw.rectangle((x1, y1, x2, y1 + 54), fill="#ecfeff", outline=accent, width=3)
        bbox = draw.textbbox((0, 0), name, font=class_font)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2, y1 + 12), name, font=class_font, fill=dark)
        current_y = y1 + 82
        for item in items:
            draw.text((x1 + 20, current_y), item, font=body_font, fill="#4b5563")
            current_y += 34

    for start, end in [
        ((420, 360), (520, 360)),
        ((880, 360), (980, 360)),
        ((700, 540), (700, 620)),
        ((1140, 500), (900, 620)),
    ]:
        draw_arrow(draw, start, end, accent)

    save_image(image, FIGURE_FILES["3.5"])


def create_sequence_diagram():
    image = Image.new("RGB", (1600, 950), "#fffefc")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    head_font = load_font(20, bold=True)
    body_font = load_font(16)
    accent = "#b45309"
    dark = "#5b3a00"
    draw.text((60, 40), "Sequence Diagram of HeartGuard Prediction Flow", font=title_font, fill=dark)

    actors = ["User", "Frontend", "FastAPI", "Risk Engine", "History"]
    x_positions = [140, 420, 700, 980, 1260]

    for actor, x in zip(actors, x_positions):
        draw.rounded_rectangle((x - 80, 120, x + 80, 180), radius=12, fill="#fff7ed", outline=accent, width=3)
        bbox = draw.textbbox((0, 0), actor, font=head_font)
        draw.text((x - (bbox[2] - bbox[0]) / 2, 138), actor, font=head_font, fill=dark)
        draw.line((x, 180, x, 850), fill="#d6d3d1", width=2)

    messages = [
        (140, 420, 250, "enter values"),
        (420, 700, 330, "POST /predict"),
        (700, 980, 410, "sanitize + score"),
        (980, 1260, 500, "save record"),
        (1260, 980, 570, "saved"),
        (980, 700, 640, "risk + confidence"),
        (700, 420, 720, "JSON response"),
        (420, 140, 800, "show result"),
    ]
    for x1, x2, y, label in messages:
        draw_arrow(draw, (x1, y), (x2, y), accent)
        draw.text((min(x1, x2) + 18, y - 28), label, font=body_font, fill=dark)

    save_image(image, FIGURE_FILES["3.6"])


def create_prediction_pipeline():
    image = Image.new("RGB", (1680, 980), "#f8fbff")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    head_font = load_font(22, bold=True)
    body_font = load_font(18)
    accent = "#1d4ed8"
    dark = "#16324f"
    muted = "#4b5563"
    draw.text((60, 40), "Detailed Prediction and Persistence Pipeline", font=title_font, fill=dark)

    lanes = [
        ("Input Layer", (80, 170, 420, 860), "#eff6ff"),
        ("API Logic", (470, 170, 850, 860), "#f8fafc"),
        ("Persistence", (900, 170, 1220, 860), "#fefce8"),
        ("Reuse Modules", (1270, 170, 1600, 860), "#f0fdf4"),
    ]
    for label, box, fill in lanes:
        draw.rounded_rectangle(box, radius=20, fill=fill, outline=accent, width=3)
        bbox = draw.textbbox((0, 0), label, font=head_font)
        draw.text((box[0] + ((box[2] - box[0]) - (bbox[2] - bbox[0])) / 2, box[1] + 18), label, font=head_font, fill=dark)

    left_steps = [
        ((120, 270, 380, 350), ("Patient form", "age, BMI, BP,\nglucose, EF,\ncreatinine")),
        ((120, 430, 380, 510), ("Client checks", "need >= 2 values,\nshow loading/status")),
        ((120, 590, 380, 670), ("Fallback path", "estimate locally if\nbackend unavailable")),
    ]
    mid_steps = [
        ((520, 240, 800, 320), ("sanitize_payload()", "drop blanks,\nkeep supported fields")),
        ((520, 390, 800, 470), ("calculate_prediction()", "weighted factor scoring,\ncoverage-based confidence")),
        ((520, 540, 800, 620), ("build_prediction_record()", "timestamp, risk, label,\nfield counts")),
        ((520, 690, 800, 770), ("POST /predict response", "risk %, confidence,\nlabel, breakdown")),
    ]
    persist_steps = [
        ((940, 300, 1180, 380), ("Excel workbook", "Predictions sheet\nwith reusable history")),
        ((940, 500, 1180, 580), ("GET /predictions", "normalize rows for\nsubsequent pages")),
    ]
    support_steps = [
        ((1310, 250, 1560, 330), ("Dashboard", "charts, streaks,\nlatest profile")),
        ((1310, 420, 1560, 500), ("Diet page", "meal guidance from\nrecent prediction")),
        ((1310, 590, 1560, 670), ("Monitoring + mood", "local logs aligned\nwith risk context")),
        ((1310, 760, 1560, 840), ("Chatbot", "prevention tips and\nheart-health Q&A")),
    ]

    for box, text in left_steps + mid_steps + persist_steps + support_steps:
        draw_box(draw, box, text, head_font, body_font, "#ffffff", accent, dark, muted, radius=18)

    arrows = [
        ((380, 310), (520, 280)),
        ((380, 470), (520, 430)),
        ((380, 630), (520, 730)),
        ((800, 280), (940, 340)),
        ((800, 430), (940, 340)),
        ((800, 580), (940, 540)),
        ((1180, 340), (1310, 290)),
        ((1180, 540), (1310, 460)),
        ((1180, 540), (1310, 630)),
        ((1180, 540), (1310, 800)),
    ]
    for start, end in arrows:
        draw_arrow(draw, start, end, accent)

    save_image(image, FIGURE_FILES["3.7"])


def create_runtime_topology():
    image = Image.new("RGB", (1680, 940), "#fffdfa")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    head_font = load_font(22, bold=True)
    body_font = load_font(18)
    accent = "#b45309"
    dark = "#5b3a00"
    muted = "#57534e"
    draw.text((60, 40), "Detailed Runtime and Deployment Topology", font=title_font, fill=dark)

    draw_box(draw, (100, 210, 390, 620), ("User Browser", "index, predict,\ndashboard, diet,\nchatbot, monitoring"), head_font, body_font, "#ffffff", accent, dark, muted)
    draw_box(draw, (520, 150, 860, 350), ("Static Frontend", "HTML + CSS + JS\nserved with app"), head_font, body_font, "#fff7ed", accent, dark, muted)
    draw_box(draw, (520, 430, 860, 760), ("FastAPI Backend", "/api/health\n/predict\n/predictions\n/chat"), head_font, body_font, "#ffffff", accent, dark, muted)
    draw_box(draw, (980, 180, 1320, 390), ("Risk Logic", "FIELD_SPECS,\nweighted scoring,\nrisk labels"), head_font, body_font, "#fefce8", accent, dark, muted)
    draw_box(draw, (980, 500, 1320, 710), ("History Storage", "predictions_history.xlsx\nOpenPyXL read/write"), head_font, body_font, "#f0fdf4", accent, dark, muted)
    draw_box(draw, (1370, 240, 1600, 360), ("Notebook Assets", ".ipynb training,\n.pkl models,\ndataset files"), head_font, body_font, "#eff6ff", accent, dark, muted)
    draw_box(draw, (1370, 540, 1600, 690), ("Deployment Layer", "render.yaml\nuvicorn backend.main:app"), head_font, body_font, "#faf5ff", accent, dark, muted)

    for start, end in [
        ((390, 300), (520, 250)),
        ((390, 520), (520, 560)),
        ((690, 350), (690, 430)),
        ((860, 560), (980, 280)),
        ((860, 620), (980, 600)),
        ((1320, 280), (1370, 300)),
        ((1320, 620), (1370, 620)),
    ]:
        draw_arrow(draw, start, end, accent)

    note_box = (170, 760, 1510, 870)
    draw.rounded_rectangle(note_box, radius=18, fill="#fff7ed", outline=accent, width=3)
    note = "Topology insight: the live web prototype uses explainable scoring and workbook persistence, while the same repository preserves richer research assets for future model-serving upgrades."
    draw.multiline_text((220, 795), note, font=body_font, fill=dark, spacing=8)

    save_image(image, FIGURE_FILES["3.8"])


def create_figures():
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    create_process_model()
    create_gantt_chart()
    create_architecture()
    create_methodology_flow()
    create_use_case()
    create_activity()
    create_class_diagram()
    create_sequence_diagram()
    create_prediction_pipeline()
    create_runtime_topology()


def add_manual_lists(doc: Document):
    add_paragraph(doc, "LIST OF FIGURES", font="Times New Roman", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10, spacing=1.0)
    figure_entries = [
        "Figure 2.1 Project Process Modelling",
        "Figure 2.2 Gantt Chart",
        "Figure 3.1 System Architecture of HeartGuard",
        "Figure 3.2 Methodology Flow Diagram of HeartGuard System",
        "Figure 3.3 Use Case Diagram of HeartGuard System",
        "Figure 3.4 Activity Diagram of HeartGuard System",
        "Figure 3.5 Class Diagram of HeartGuard System",
        "Figure 3.6 Sequence Diagram of HeartGuard System",
        "Figure 3.7 Detailed Prediction and Persistence Pipeline",
        "Figure 3.8 Detailed Runtime and Deployment Topology",
    ]
    for entry in figure_entries:
        add_paragraph(doc, entry, align=WD_ALIGN_PARAGRAPH.LEFT, after=2, spacing=1.0)
    doc.add_page_break()

    add_paragraph(doc, "LIST OF TABLES", font="Times New Roman", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10, spacing=1.0)
    table_entries = [
        "Table 1.1 Literature Survey",
        "Table 2.1 Requirement Summary",
        "Table 2.2 Feasibility Snapshot",
        "Table 2.3 Project Scheduling",
        "Table 3.1 Module Description of HeartGuard System",
        "Table 3.3 Design Goal to Solution Mapping",
        "Table 4.1 Backend Endpoints and Responsibilities",
        "Table 4.2 Feature to Component Mapping",
        "Table 5.1 Testing Summary",
        "Table 6.1 Results and Discussion Summary",
        "Table A.1 Dataset and Variable Summary",
    ]
    for entry in table_entries:
        add_paragraph(doc, entry, align=WD_ALIGN_PARAGRAPH.LEFT, after=2, spacing=1.0)
    doc.add_page_break()

    add_paragraph(doc, "ABBREVIATIONS", font="Times New Roman", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10, spacing=1.0)
    abbreviations = [
        "HG - HeartGuard",
        "HF - Heart Failure",
        "AI - Artificial Intelligence",
        "ML - Machine Learning",
        "API - Application Programming Interface",
        "EHR - Electronic Health Record",
        "ROC - Receiver Operating Characteristic",
        "AUC - Area Under Curve",
        "EF - Ejection Fraction",
        "BMI - Body Mass Index",
    ]
    for item in abbreviations:
        add_paragraph(doc, item, align=WD_ALIGN_PARAGRAPH.LEFT, after=2, spacing=1.0)
    doc.add_page_break()

    add_paragraph(doc, "TABLE OF CONTENTS", font="Times New Roman", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10, spacing=1.0)
    toc_lines = [
        "Chapter 1 Introduction",
        "1.1 Background and Basics",
        "1.2 Literature Survey and Research Gap",
        "1.3 Project Undertaken",
        "1.4 Problem Definition",
        "1.5 Scope Statement",
        "1.6 Organization of Project Report",
        "Chapter 2 Project Planning and Management",
        "Chapter 3 Analysis and Design",
        "Chapter 4 Implementation and Coding",
        "Chapter 5 Testing",
        "Chapter 6 Results and Discussion",
        "Chapter 7 Conclusion",
        "Chapter 8 Future Work",
        "References",
        "Annexure",
    ]
    for line in toc_lines:
        add_paragraph(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, after=2, spacing=1.0)
    doc.add_page_break()


def build_report():
    create_figures()

    dataset_counts = {
        "master_rows": count_csv_rows(MASTER_DATASET),
        "master_cols": count_csv_columns(MASTER_DATASET),
        "cardio_rows": count_csv_rows(CARDIO_DATASET),
        "cardio_cols": count_csv_columns(CARDIO_DATASET),
        "diabetes_rows": count_csv_rows(DIABETES_DATASET),
        "diabetes_cols": count_csv_columns(DIABETES_DATASET),
        "framingham_rows": count_csv_rows(FRAMINGHAM_DATASET),
        "framingham_cols": count_csv_columns(FRAMINGHAM_DATASET),
        "heart_failure_rows": count_csv_rows(HEART_FAILURE_DATASET),
        "heart_failure_cols": count_csv_columns(HEART_FAILURE_DATASET),
    }
    accuracy, auc = read_training_metrics()
    backend_code = read_excerpt(BACKEND_FILE, "class PatientData(BaseModel):", '@app.get("/api/health")', max_lines=72)
    history_code = read_excerpt(BACKEND_FILE, "def ensure_predictions_workbook() -> None:", "class PatientData(BaseModel):", max_lines=68)
    predict_script_code = read_excerpt(PREDICT_SCRIPT_FILE, "async function handlePrediction(event) {", "function sanitizePercentage(value) {", max_lines=74)
    dashboard_code = read_excerpt(DASHBOARD_SCRIPT_FILE, "async function loadPredictionHistory() {", "function renderStats(statsData) {", max_lines=78)
    chatbot_code = read_excerpt(BACKEND_FILE, "class ChatRequest(BaseModel):", "def answer_stress(_: str) -> str:", max_lines=88)
    diet_code = read_excerpt(DIET_SCRIPT_FILE, "async function initializeDietPage() {", "function renderDietRecommendation(profile, prediction) {", max_lines=92)
    notebook_code = read_notebook_training_excerpt(TRAINING_NOTEBOOK)

    doc = Document()
    set_page_layout(doc)
    set_default_style(doc)

    # Cover page
    add_paragraph(doc, "", before=90, after=0, spacing=1.0)
    add_center_block(
        doc,
        [
            ("A PROJECT REPORT ON", 14, True),
            (REPORT_TITLE, 16, True),
            ("SUBMITTED TO THE SAVITRIBAI PHULE PUNE UNIVERSITY, PUNE", 12, False),
            ("IN PARTIAL FULFILLMENT OF THE REQUIREMENTS", 12, False),
            ("FOR THE AWARD OF THE DEGREE", 12, False),
            ("BACHELOR OF ENGINEERING", 14, True),
            ("IN", 12, False),
            ("COMPUTER ENGINEERING", 14, True),
            ("OF", 12, False),
            ("SAVITRIBAI PHULE PUNE UNIVERSITY", 12, True),
            ("BY", 12, False),
        ],
    )
    for name, roll in TEAM_MEMBERS:
        add_paragraph(doc, f"{name}    Roll No: {roll}", font="Times New Roman", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2, spacing=1.0)
    add_paragraph(doc, "", after=4, spacing=1.0)
    add_center_block(
        doc,
        [
            ("Under the guidance of", 12, False),
            ("Prof. N. G. BHOJNE", 13, True),
            ("DEPARTMENT OF COMPUTER ENGINEERING", 12, True),
            ("SINHGAD COLLEGE OF ENGINEERING, PUNE-41", 12, True),
            ("Accredited by NAAC with A+ Grade", 11, False),
            ("2025-26", 12, True),
        ],
    )
    doc.add_page_break()

    # Certificate
    add_center_block(
        doc,
        [
            ("Sinhgad Technical Education Society", 12, False),
            ("Sinhgad College of Engineering, Pune-41", 12, False),
            ("Department of Computer Engineering", 12, False),
        ],
    )
    add_paragraph(doc, "Date:", align=WD_ALIGN_PARAGRAPH.LEFT, after=8, spacing=1.0)
    add_paragraph(doc, "CERTIFICATE", font="Times New Roman", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12, spacing=1.0)
    certificate_text = (
        f'This is to certify that the project report entitled "{REPORT_TITLE}" submitted by '
        "Harsh Bandal, Pankaj Aswale, Samruddhi Bharde, and Sonali Bhojane is a bonafide work carried out under "
        "the supervision of Prof. N. G. Bhojne and is approved for the partial fulfillment of the requirements "
        "of Savitribai Phule Pune University, Pune for the award of the degree of Bachelor of Engineering "
        "(Computer Engineering) during the year 2025-26."
    )
    add_paragraph(doc, certificate_text, after=10)
    signature_table = doc.add_table(rows=2, cols=3)
    signature_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    signature_table.style = "Table Grid"
    set_column_widths(signature_table, [Inches(2.1), Inches(2.1), Inches(2.1)])
    configure_cell(signature_table.cell(0, 0), "Prof. N. G. Bhojne\nGuide\nDepartment of Computer Engineering", align=WD_ALIGN_PARAGRAPH.CENTER)
    configure_cell(signature_table.cell(0, 1), "Prof. R. H. Borhade\nHead\nDepartment of Computer Engineering", align=WD_ALIGN_PARAGRAPH.CENTER)
    configure_cell(signature_table.cell(0, 2), "Dr. S. D. Lokhande\nPrincipal\nSinhgad College of Engineering", align=WD_ALIGN_PARAGRAPH.CENTER)
    configure_cell(signature_table.cell(1, 0), "Harsh Bandal\n405A101", align=WD_ALIGN_PARAGRAPH.CENTER)
    configure_cell(signature_table.cell(1, 1), "Pankaj Aswale\n405A096", align=WD_ALIGN_PARAGRAPH.CENTER)
    configure_cell(signature_table.cell(1, 2), "Samruddhi Bharde / Sonali Bhojane\n405A092 / 405A090", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # Acknowledgement
    add_paragraph(doc, "ACKNOWLEDGEMENT", font="Times New Roman", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12, spacing=1.0)
    acknowledgement = (
        "We express our sincere gratitude to Prof. N. G. Bhojne for his constant guidance, motivation, and technical insight throughout "
        "the development of HeartGuard. His support helped us move from an initial predictive-healthcare idea to a practical web-based project "
        "that combines forecasting, persistence, and patient-support modules.\n\n"
        "We are also thankful to Dr. S. D. Lokhande, Principal of Sinhgad College of Engineering, Pune, and Prof. R. H. Borhade, Head of the "
        "Department of Computer Engineering, for providing the academic environment, infrastructure, and encouragement necessary for this work. "
        "We further acknowledge the contribution of our department staff, classmates, and the openly available healthcare datasets and tools that "
        "made experimentation, validation, and report preparation possible."
    )
    for block in acknowledgement.split("\n\n"):
        add_paragraph(doc, block)
    add_paragraph(doc, "Harsh Bandal\nPankaj Aswale\nSamruddhi Bharde\nSonali Bhojane", align=WD_ALIGN_PARAGRAPH.RIGHT, after=2, spacing=1.0)
    doc.add_page_break()

    # Abstract
    add_paragraph(doc, "ABSTRACT", font="Times New Roman", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=12, spacing=1.0)
    abstract = (
        "Heart failure remains one of the leading causes of repeated hospitalization, and timely identification of high-risk patients is still "
        "a major challenge in preventive healthcare. HeartGuard is proposed as a combined forecasting and support platform that studies heart-failure-"
        "aware risk prediction while also providing a practical web workflow for continuity after assessment. The project uses a FastAPI backend, "
        "modular HTML/CSS/JavaScript frontend pages, persistent Excel-backed prediction history, and a multi-source clinical data foundation that "
        f"includes {dataset_counts['master_rows']:,} harmonized records across cardiovascular, diabetes, Framingham, and heart-failure datasets. "
        "The current public prototype emphasizes explainability and sparse-input usability through a weighted scoring engine that accepts incomplete "
        "clinical inputs, estimates confidence, stores results, and connects them to dashboard review, diet planning, AI assistant support, mood logging, "
        "medication tracking, and vital monitoring.\n\n"
        f"Training artifacts in the project repository demonstrate exploratory supervised-model performance with accuracy {accuracy} and ROC-AUC {auc}, "
        "while the deployed application intentionally prioritizes interpretability and robustness over black-box complexity. Functional validation confirms "
        "partial-input prediction, dashboard continuity after refresh, and successful handoff from forecasting to supportive follow-up modules. The project "
        "therefore presents HeartGuard as a realistic academic prototype for preventive insight generation, readmission-aware monitoring, and future model-integrated healthcare analytics."
    )
    for block in abstract.split("\n\n"):
        add_paragraph(doc, block)
    doc.add_page_break()

    add_manual_lists(doc)

    # Chapter 1
    add_chapter_title(doc, "1", "INTRODUCTION")
    add_section_title(doc, "1.1 Background and Basics")
    background_paragraphs = [
        "Heart disease has emerged as one of the most significant health challenges of the modern era. Within the broader cardiovascular domain, heart failure is especially concerning because it is chronic, progressive, and closely associated with repeated admissions, long recovery cycles, and high treatment costs. Even after discharge, many patients require careful follow-up because symptoms such as fatigue, breathlessness, fluid retention, and unstable blood pressure can deteriorate quickly when not monitored properly.",
        "Hospital readmission after heart-failure treatment remains a serious burden for patients, families, and healthcare systems. A large share of these readmissions are influenced by delayed recognition of risk, weak continuity between discharge and home care, and limited preventive support once a patient leaves the hospital. This makes early forecasting valuable not only as a prediction problem, but also as a clinical workflow problem where insight must lead to timely action.",
        "Traditional readmission assessment methods often depend on a narrow set of variables or linear statistical assumptions. However, real patient outcomes are shaped by interacting factors such as age, metabolic status, blood pressure, cholesterol, glucose, smoking behavior, kidney function, and ventricular performance. As a result, there is growing interest in systems that can combine predictive modeling with explainable preventive recommendations."
    ]
    for p in background_paragraphs:
        add_paragraph(doc, p)

    add_section_title(doc, "1.2 Literature Survey and Research Gap")
    add_paragraph(doc, "Recent research has shown that machine learning can improve outcome forecasting in cardiovascular care, especially for mortality and readmission problems. At the same time, published studies also reveal important limitations related to interpretability, generalization, sparse data, and lack of clinical-action support.", after=6)
    add_paragraph(doc, "Table 1.1 Literature Survey", font="Times New Roman", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, spacing=1.0)
    add_table(
        doc,
        ["Study", "Approach", "Finding", "Limitation"],
        [
            ["Hajishah et al. (2025)", "Meta-analysis of machine learning for heart failure mortality and readmission", "Confirmed predictive potential across model families", "Interpretability remains limited in many high-performing models"],
            ["Comparative review (2023)", "30-day HF readmission models", "AUC varied significantly across datasets and methods", "No consistent feature set or deployment standard"],
            ["Ru et al. (2023)", "RF, XGBoost, deep learning on HFrEF readmissions", "Moderate performance for short-term forecasting", "Data imbalance and missing values reduced reliability"],
            ["Graph-based study (2024)", "Clinical similarity graph model", "Improved readmission discrimination in one setting", "Single-center training reduced generalizability"],
        ],
        [Inches(1.55), Inches(1.7), Inches(1.55), Inches(1.6)],
    )
    add_paragraph(doc, "The review suggests a clear research gap. Many systems concentrate on prediction quality alone, but fewer connect prediction with practical preventive follow-up. Likewise, models that perform well in isolated experiments are often difficult to explain to clinicians or use when patient inputs are incomplete. HeartGuard addresses this gap by combining data-driven experimentation with an explainable web prototype that supports continuity after prediction.")

    add_section_title(doc, "1.3 Project Undertaken")
    project_paragraphs = [
        "The project undertaken in this report is the design and implementation of HeartGuard, a preventive insight platform for heart-failure readmission forecasting. Rather than treating risk estimation as a one-time classification event, HeartGuard connects prediction with dashboard continuity, personalized diet guidance, AI-assisted health information, mood tracking, medication management, and vital-sign monitoring.",
        "This dual structure is important because the current repository contains both training-oriented machine learning assets and a deployed end-user prototype. Notebook experiments and stored model artifacts support the forecasting research dimension, while the FastAPI and frontend stack deliver an interpretable and usable experience for real-world demonstration and academic evaluation.",
    ]
    for p in project_paragraphs:
        add_paragraph(doc, p)

    add_section_title(doc, "1.4 Problem Definition")
    add_paragraph(doc, "The core problem addressed by HeartGuard is the difficulty of identifying patients at higher risk of post-discharge deterioration or readmission early enough to enable preventive intervention. Existing approaches often fail when data are incomplete, difficult to interpret, or disconnected from meaningful follow-up tools. The project therefore seeks to create a system that can estimate heart-related risk from limited but clinically useful inputs and convert that estimate into practical support for patients and healthcare stakeholders.")

    add_section_title(doc, "1.5 Scope Statement")
    add_paragraph(doc, "The scope of the project includes multi-source dataset preparation, exploratory machine learning training, implementation of an explainable weighted scoring predictor, creation of a browser-based FastAPI application, persistence of prediction history, and integration of supportive wellness modules around the prediction workflow. The current scope does not include real-time hospital EHR integration, production-grade multi-user authentication, or certified clinical deployment. Those areas are reserved for future work.")

    add_section_title(doc, "1.6 Organization of Project Report")
    add_bullets(
        doc,
        [
            "Chapter 1 introduces the domain, problem, literature context, and project scope.",
            "Chapter 2 explains project planning, requirements, feasibility, process modelling, and scheduling.",
            "Chapter 3 presents analysis and design, including architecture, module breakdown, workflow, and UML-style representations.",
            "Chapter 4 describes the implementation of the frontend, backend, persistence, and exploratory model assets.",
            "Chapter 5 discusses unit, integration, system, and functional testing.",
            "Chapter 6 presents results and implementation-level discussion.",
            "Chapter 7 concludes the report and Chapter 8 outlines future work.",
        ],
    )
    doc.add_page_break()

    # Chapter 2
    add_chapter_title(doc, "2", "PROJECT PLANNING AND MANAGEMENT")
    add_section_title(doc, "2.1 Introduction")
    add_paragraph(doc, "Project planning and management were essential in translating HeartGuard from an abstract predictive-healthcare idea into a usable academic prototype. The work involved dataset preparation, exploratory notebook-based training, frontend and backend implementation, persistence design, and reporting. A structured process reduced ambiguity and made it easier to align technical effort with the final project objective.")

    add_section_title(doc, "2.2 Software Requirement Specification (SRS)")
    add_subsection_title(doc, "2.2.1 System Overview")
    add_paragraph(doc, "HeartGuard is a web-based health-support system with a FastAPI backend and multiple frontend pages. The system accepts patient health values, estimates cardiovascular and heart-failure-aware risk, stores the prediction history, and reuses the latest profile across support modules such as dashboard summaries, diet planning, mood tracking, medication support, AI guidance, and vitals monitoring.")
    add_subsection_title(doc, "2.2.2 Functional Requirements")
    add_paragraph(doc, "Table 2.1 Requirement Summary", font="Times New Roman", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, spacing=1.0)
    add_table(
        doc,
        ["Category", "Requirement", "Current project evidence", "Expected benefit"],
        [
            ["Input", "Accept at least two clinical values for prediction", "Backend validates sparse payloads and scores supported fields", "Improves usability when users lack complete records"],
            ["Prediction", "Return risk, confidence, and label", "FastAPI /predict endpoint and frontend result card", "Makes outputs interpretable for review"],
            ["Persistence", "Store prediction history", "Excel-backed workbook with timestamped entries", "Supports continuity after refresh"],
            ["Support workflow", "Reuse latest prediction in other modules", "Dashboard and diet pages load most recent saved history", "Connects prediction to preventive action"],
            ["Health guidance", "Answer general heart-health questions", "Rule-based chatbot endpoint and UI", "Improves accessibility of educational support"],
        ],
        [Inches(1.0), Inches(1.65), Inches(2.2), Inches(1.6)],
    )
    add_subsection_title(doc, "2.2.3 Non-Functional Requirements")
    add_bullets(
        doc,
        [
            "Usability: the interface should remain understandable for non-technical users and support module-to-module navigation.",
            "Reliability: risk calculation, history recovery, and local fallback behavior should function consistently.",
            "Maintainability: the backend should allow replacement of Excel persistence with a database later.",
            "Scalability: the modular frontend and API structure should permit additional support pages and future model-serving integration.",
            "Confidentiality: patient information must be treated as sensitive even in prototype form.",
        ],
    )
    add_subsection_title(doc, "2.2.4 Deployment Environment")
    add_paragraph(doc, "HeartGuard currently targets Windows-based local execution and free cloud deployment through Render. The implementation stack uses Python with FastAPI, Uvicorn, OpenPyXL, and Pydantic on the backend, while the frontend uses HTML, CSS, JavaScript, browser localStorage, and Chart.js for selected dashboards.")
    add_subsection_title(doc, "2.2.5 External Interface Requirements")
    add_paragraph(doc, "The user-facing interface consists of browser pages for landing, login, prediction, dashboard, AI assistant, diet plan, monitoring, mood logging, medications, breathing, community, quiz, and SOS support. The backend interface consists of HTTP endpoints for health, prediction, history, and chatbot operations.")

    add_section_title(doc, "2.3 Feasibility and Resource Planning")
    add_paragraph(doc, "Feasibility was evaluated from technical, operational, schedule, and economic perspectives. The project was technically feasible because the required datasets, Python libraries, and frontend tooling were all accessible. It was operationally feasible because the chosen prototype design reduced complexity by using workbook persistence instead of a full hospital database.")
    add_paragraph(doc, "Table 2.2 Feasibility Snapshot", font="Times New Roman", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, spacing=1.0)
    add_table(
        doc,
        ["Dimension", "Observation", "Project implication"],
        [
            ["Technical", "Existing datasets, notebooks, and model artifacts were available", "Supported rapid experimentation and prototype alignment"],
            ["Operational", "Users can interact through simple browser pages", "Reduced training burden for demonstration use"],
            ["Economic", "Open-source stack and free deployment option are available", "Kept project cost low"],
            ["Schedule", "Work could be split into data, backend, frontend, and reporting phases", "Made semester delivery realistic"],
        ],
        [Inches(1.2), Inches(2.8), Inches(2.4)],
    )

    add_section_title(doc, "2.4 Project Process Modelling")
    add_paragraph(doc, "The project followed a structured waterfall-like progression because the main milestones were well defined: understand the domain, prepare data, explore models, implement the system, validate the workflow, and compile the final report.")
    add_figure(doc, FIGURE_FILES["2.1"], "Figure 2.1 Project Process Modelling", 6.4)

    add_section_title(doc, "2.5 Project Scheduling")
    add_paragraph(doc, "The semester schedule balanced domain study, coding effort, testing, and report writing. Since the work included both forecasting research and user-facing implementation, scheduling had to accommodate iterative refinement between backend logic and frontend usability.")
    add_paragraph(doc, "Table 2.3 Project Scheduling", font="Times New Roman", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, spacing=1.0)
    add_table(
        doc,
        ["Phase", "Main activity", "Indicative period"],
        [
            ["Phase I", "Problem study, literature review, requirement definition", "August-September"],
            ["Phase II", "Dataset preparation and notebook experiments", "September-October"],
            ["Phase III", "Backend and frontend implementation", "October-November"],
            ["Phase IV", "Testing, refinement, and report preparation", "November-January"],
        ],
        [Inches(1.0), Inches(4.0), Inches(1.6)],
    )
    add_figure(doc, FIGURE_FILES["2.2"], "Figure 2.2 Gantt Chart", 6.45)
    doc.add_page_break()

    # Chapter 3
    add_chapter_title(doc, "3", "ANALYSIS AND DESIGN")
    add_section_title(doc, "3.1 System Analysis")
    add_paragraph(doc, "System analysis for HeartGuard focused on how to translate a predictive-healthcare problem into a practical academic application. The analysis phase defined the required inputs, the processing pipeline, the relationship between the prediction engine and support modules, and the data continuity needed for meaningful review after the initial estimate.")
    add_bullets(
        doc,
        [
            "Input: clinical and lifestyle values such as age, BMI, blood pressure, glucose, cholesterol, smoking status, diabetes status, heart rate, ejection fraction, and serum creatinine.",
            "Process: payload sanitization, weighted scoring or research-model analysis, confidence estimation, persistence, and module reuse.",
            "Output: risk percentage, confidence value, risk label, saved history, and follow-up guidance.",
        ],
    )

    add_section_title(doc, "3.2 System Architecture")
    add_paragraph(doc, "The architecture combines data assets, an API layer, workbook persistence, browser interfaces, and health-support modules. Importantly, the repository includes exploratory model artifacts alongside the deployed prototype, so the design must support both predictive experimentation and user-facing demonstration.")
    add_figure(doc, FIGURE_FILES["3.1"], "Figure 3.1 System Architecture of HeartGuard", 6.4)

    add_section_title(doc, "3.3 Data and Prediction Flow")
    add_paragraph(doc, "The methodology begins with requirement analysis and dataset alignment, proceeds through exploratory modeling, and ends with an explainable deployment path. This is consistent with the present codebase, where notebook results inform feature relevance while the live application exposes a more transparent decision flow.")
    add_figure(doc, FIGURE_FILES["3.2"], "Figure 3.2 Methodology Flow Diagram of HeartGuard System", 6.4)

    add_section_title(doc, "3.4 Module Description")
    add_paragraph(doc, "Table 3.1 Module Description of HeartGuard System", font="Times New Roman", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, spacing=1.0)
    add_table(
        doc,
        ["Module", "Primary responsibility", "Current implementation evidence"],
        [
            ["Prediction", "Accept user inputs and estimate risk", "predict.html, predict-script.js, backend /predict endpoint"],
            ["Dashboard", "Recover latest history and summarize trends", "dashboard.html, dashboard-script.js, /predictions endpoint"],
            ["Diet Support", "Generate meal guidance from latest profile", "diet.html, diet-script.js"],
            ["Chatbot", "Provide heart-health education and support", "chatbot.html, /chat endpoint"],
            ["Mood / Monitoring / Medications", "Track supportive wellness data locally", "mood-script.js, monitoring-script.js, medications-script.js"],
            ["Community / Quiz / SOS", "Broaden engagement and awareness", "community.html, quiz.html, sos.html"],
        ],
        [Inches(1.4), Inches(2.0), Inches(3.0)],
    )

    add_section_title(doc, "3.5 UML and Interaction Design")
    add_paragraph(doc, "The following diagrams summarize how users interact with HeartGuard and how the system components cooperate to turn health inputs into saved preventive insights.")
    add_figure(doc, FIGURE_FILES["3.3"], "Figure 3.3 Use Case Diagram of HeartGuard System", 6.4)
    add_figure(doc, FIGURE_FILES["3.4"], "Figure 3.4 Activity Diagram of HeartGuard System", 6.25)
    add_figure(doc, FIGURE_FILES["3.5"], "Figure 3.5 Class Diagram of HeartGuard System", 6.4)
    add_figure(doc, FIGURE_FILES["3.6"], "Figure 3.6 Sequence Diagram of HeartGuard System", 6.4)
    add_paragraph(doc, "Because the reference report includes more implementation-oriented visual detail, HeartGuard also documents the internal runtime flow more explicitly. The next two diagrams connect the user journey, API operations, workbook persistence, and downstream wellness modules at a finer level of granularity.")
    add_figure(doc, FIGURE_FILES["3.7"], "Figure 3.7 Detailed Prediction and Persistence Pipeline", 6.45)
    add_figure(doc, FIGURE_FILES["3.8"], "Figure 3.8 Detailed Runtime and Deployment Topology", 6.45)

    add_section_title(doc, "3.6 Mathematical and Scoring Model")
    add_paragraph(doc, "The deployed predictor uses a weighted scoring model. Each supported field contributes a clinically motivated factor score and a corresponding weight. The aggregate risk percentage is computed from the weighted average of all supplied variables, and a confidence score is then estimated from feature coverage and signal strength. This approach offers two advantages: the interface remains usable when only partial data are available, and the result is easier to explain than a hidden black-box probability.")
    add_paragraph(doc, "In parallel, the repository also stores notebook-based supervised learning experiments and serialized model artifacts for datasets such as Framingham, heart failure, cardio, diabetes, and a master dataset. This makes the design extensible: the present workflow is explainable and educational, while future versions can reconnect the web interface to more advanced validated model-serving pipelines.")
    add_section_title(doc, "3.7 Design Considerations")
    add_paragraph(doc, "Several design decisions shaped the final HeartGuard prototype. First, the project chooses explainability over hidden complexity at deployment time, which is why weighted scoring remains the live interface while notebook-based models remain available as research assets. Second, continuity was treated as a core requirement rather than an optional enhancement, leading to workbook history, dashboard recovery, and module reuse from the latest prediction.")
    add_paragraph(doc, "Third, the user experience was intentionally broadened beyond a single prediction page. The surrounding modules for diet, monitoring, mood, medications, community support, and chatbot assistance reflect the idea that readmission-aware care should connect risk insight to daily preventive action. Finally, the design keeps its deployment limitations visible, particularly around workbook persistence and the lack of role-based security, so that the system is interpreted honestly during academic review.")
    add_bullets(
        doc,
        [
            "Modularity enables future migration from workbook persistence to a relational database.",
            "API-based separation allows the frontend to evolve independently from model-serving changes.",
            "Local support modules demonstrate continuity even before clinical integration is available.",
        ],
    )
    add_section_title(doc, "3.8 Requirements Traceability")
    add_paragraph(doc, "Table 3.3 links key design expectations to the architectural elements that were selected in the final prototype. This traceability view helps academic reviewers see that the observed implementation choices are not arbitrary additions, but responses to concrete project needs.")
    add_paragraph(doc, "Table 3.3 Design Goal to Solution Mapping", font="Times New Roman", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, spacing=1.0)
    add_table(
        doc,
        ["Design goal", "Observed need", "Implemented response"],
        [
            ["Early risk visibility", "Patients need a quick, understandable indication of concern", "Weighted scoring model returns risk, label, confidence, and explanation-ready inputs"],
            ["Continuity of care", "Users should not lose context after one prediction", "Workbook-backed history and dashboard recovery preserve recent assessments"],
            ["Preventive actionability", "Prediction alone is insufficient for lifestyle support", "Diet, monitoring, mood, medications, and chatbot modules extend post-result engagement"],
            ["Low-cost deployment", "Prototype should run on modest academic infrastructure", "FastAPI, static frontend, and minimal dependencies reduce operational overhead"],
        ],
        [Inches(1.6), Inches(2.0), Inches(2.9)],
    )
    doc.add_page_break()

    # Chapter 4
    add_chapter_title(doc, "4", "IMPLEMENTATION AND CODING")
    add_section_title(doc, "4.1 Introduction")
    add_paragraph(doc, "This chapter documents how the HeartGuard design was translated into executable software. In line with the common reference report, the implementation discussion moves from architectural components to source-code listings so that the reader can connect design decisions with actual program structure.")
    add_paragraph(doc, "HeartGuard is implemented as a browser-based application with a FastAPI backend, static frontend pages, workbook-backed prediction persistence, and client-side support modules. The deployed workflow favors explainability and continuity of care rather than opaque model serving, which is why the current codebase exposes clear scoring logic, normalized history handling, and reusable patient-facing guidance modules.")

    add_section_title(doc, "4.2 Software Components")
    add_paragraph(doc, "The frontend is implemented with modular HTML, CSS, and JavaScript pages. The landing page introduces the platform as an AI health system, while the prediction, dashboard, diet, chatbot, mood, monitoring, medications, breathing, community, quiz, and SOS pages extend the platform into a broader wellness workflow.")
    add_paragraph(doc, "The backend is built with FastAPI and a lightweight set of dependencies: FastAPI, Uvicorn, OpenPyXL, and Pydantic. The backend serves health, prediction, history, and chatbot endpoints, and it also mounts the frontend for browser access. Navigation is handled through consistent sidebar structures across the core pages, which helps users move from one support module to another without losing context.")
    add_paragraph(doc, "Table 4.1 Backend Endpoints and Responsibilities", font="Times New Roman", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, spacing=1.0)
    add_table(
        doc,
        ["Endpoint", "Method", "Purpose", "Observed implementation"],
        [
            ["/api/health", "GET", "Verify backend availability", "Returns HeartGuard API status message"],
            ["/predict", "POST", "Compute risk, confidence, label, and save record", "Sanitizes payload, scores fields, appends workbook history"],
            ["/predictions", "GET", "Load saved history", "Reads Excel workbook and returns items"],
            ["/chat", "POST", "Answer health-support prompts", "Routes questions to rule-based response handlers"],
        ],
        [Inches(1.3), Inches(0.8), Inches(2.0), Inches(2.4)],
    )

    add_section_title(doc, "4.3 Code Listing")
    add_subsection_title(doc, "4.3.1 backend/main.py")
    add_paragraph(doc, "The backend file is the operational core of HeartGuard. It defines the patient input schema, sanitizes incoming values, computes weighted risk and confidence scores, prepares records for storage, and exposes the prediction API used by the browser interface.")
    add_code_block(doc, backend_code)

    add_subsection_title(doc, "4.3.2 frontend/predict-script.js")
    add_paragraph(doc, "The prediction page script coordinates the main user journey. It validates that enough inputs are present, calls the backend API, handles failures gracefully, and stores fallback results locally when the server is not available.")
    add_code_block(doc, predict_script_code)

    add_subsection_title(doc, "4.3.3 frontend/dashboard-script.js")
    add_paragraph(doc, "The dashboard script demonstrates how persisted predictions are reused after scoring. It reloads history from the backend when possible, normalizes records into a chart-friendly form, computes streaks and summary statistics, and updates the visual dashboard state.")
    add_code_block(doc, dashboard_code)

    add_subsection_title(doc, "4.3.4 backend/main.py workbook persistence")
    add_paragraph(doc, "The following excerpt shows how HeartGuard maintains reusable prediction history. It creates the workbook structure when needed, reads saved rows for dashboard recovery, and appends fresh prediction records after each successful inference.")
    add_code_block(doc, history_code)

    add_subsection_title(doc, "4.3.5 backend/main.py chatbot logic")
    add_paragraph(doc, "HeartGuard also includes a rule-based support chatbot. This code fragment shows the request schema, route handling, keyword matching, and topic-specific response dispatch used to answer questions on diet, symptoms, blood pressure, cholesterol, exercise, and risk-reduction habits.")
    add_code_block(doc, chatbot_code)

    add_subsection_title(doc, "4.3.6 frontend/diet-script.js")
    add_paragraph(doc, "The diet module reuses the latest prediction profile to personalize meal guidance. The excerpt below shows how the page loads the latest saved prediction, normalizes stored values, and builds a calorie-aware, risk-aware dietary plan from blood pressure, glucose, cholesterol, BMI, and smoking-related signals.")
    add_code_block(doc, diet_code)

    add_subsection_title(doc, "4.3.7 notebooks/master_model_training.ipynb")
    add_paragraph(doc, "The notebook excerpt captures the research-side training workflow preserved in the repository. It demonstrates preprocessing, train-test splitting, supervised model fitting, and metric reporting that informed the broader predictive direction of the project.")
    add_code_block(doc, notebook_code)

    add_section_title(doc, "4.4 Data Persistence and Local Support Modules")
    add_paragraph(doc, "Prediction history is persisted in an Excel workbook stored under the data/final directory. This choice is intentionally lightweight and transparent for prototype use, allowing saved results to be inspected and reloaded easily. The dashboard consumes this history to avoid blank states after refresh. Several lifestyle modules such as mood tracking, medication management, and vitals monitoring store their entries in browser localStorage, which keeps the interface responsive even without a larger database.")

    add_section_title(doc, "4.5 Exploratory ML Assets and Model Artifacts")
    add_paragraph(doc, f"The repository contains notebook-based experiments and serialized model artifacts for cardio, diabetes, Framingham, heart failure, and master datasets. The master notebook records exploratory supervised-model performance of accuracy {accuracy} and ROC-AUC {auc}. These artifacts are important because they preserve the research dimension of the project, even though the deployed prototype currently prioritizes explainable weighted scoring for user-facing interaction.")
    add_paragraph(doc, "The presence of multiple .pkl files and harmonized datasets shows that the project was not limited to a single toy experiment. Instead, it explored how multi-source healthcare data might be aligned and reused for future ensemble or model-serving integration.")

    add_section_title(doc, "4.6 Deployment Readiness")
    add_paragraph(doc, "The project includes local launch scripts and a Render deployment configuration. The render.yaml file defines a free web-service deployment that installs requirements and starts Uvicorn with backend.main:app. The accompanying deployment note also records an important limitation: Excel-backed persistence is ephemeral on many free cloud platforms, so durable multi-user deployment would require migration to SQLite or PostgreSQL.")
    add_section_title(doc, "4.7 Code Organization Snapshot")
    add_paragraph(doc, "The present code organization is straightforward enough for maintenance and academic demonstration. Backend logic is concentrated in a single FastAPI application file, which makes it easy to trace prediction, history, and chatbot behavior during review. The frontend is split by page, with each major support function receiving its own HTML, CSS, and JavaScript files. This page-level separation helped the team iterate quickly while keeping features understandable during documentation and testing.")
    add_bullets(
        doc,
        [
            "Frontend page set includes landing, login, predict, dashboard, chatbot, mood, diet, monitoring, medications, community, quiz, breathing, and SOS pages.",
            "Model and dataset assets are stored separately from the deployed predictor, preserving a clean boundary between research files and runtime behavior.",
            "Deployment scripts and documentation are included directly in the repository for reproducible local and free-cloud execution.",
        ],
    )
    add_section_title(doc, "4.8 Implementation Traceability")
    add_paragraph(doc, "Table 4.2 summarizes how the most visible user-facing capabilities map to concrete files and runtime components in the project. This snapshot is useful during viva review because it allows an examiner to connect interface features with the code units responsible for them.")
    add_paragraph(doc, "Table 4.2 Feature to Component Mapping", font="Times New Roman", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, spacing=1.0)
    add_table(
        doc,
        ["Feature", "Primary implementation area", "Execution note"],
        [
            ["Risk prediction", "backend/main.py and predict.html", "Frontend submits form data to POST /predict and receives risk analytics"],
            ["Saved history dashboard", "dashboard.html, dashboard-script.js, GET /predictions", "Workbook records are normalized for charts, streaks, and summaries"],
            ["Diet personalization", "diet-script.js", "Latest prediction profile influences meal-plan and hydration suggestions"],
            ["Support chatbot", "chatbot-script.js and POST /chat", "Rule-based responses provide educational and motivational guidance"],
        ],
        [Inches(1.5), Inches(2.4), Inches(2.6)],
    )
    doc.add_page_break()

    # Chapter 5
    add_chapter_title(doc, "5", "TESTING")
    add_section_title(doc, "5.1 Unit Testing")
    add_paragraph(doc, "Unit-level validation focused on the correctness of the prediction workflow, data sanitation, dashboard normalization, diet-profile derivation, chatbot request handling, mood logging, medication toggling, and vitals chart preparation. These checks ensured that the application components behaved predictably before they were evaluated as a complete workflow.")

    add_section_title(doc, "5.2 Integration Testing")
    add_paragraph(doc, "Integration testing centered on the interaction between frontend pages and backend APIs. The most important integration path was prediction submission, workbook persistence, dashboard recovery, and profile reuse on the diet page. Additional integration checks considered the chatbot endpoint and the fallback behavior when a backend request was unavailable.")

    add_section_title(doc, "5.3 System Testing")
    add_paragraph(doc, "System testing exercised the entire HeartGuard flow from initial input entry to result display and downstream support usage. The goal was not only to verify that a score appeared, but also to verify that the platform behaved as a coherent preventive-care workflow. This included navigation, history continuity, local wellness logging, and clarity of guidance.")

    add_section_title(doc, "5.4 Validation and Observations")
    add_paragraph(doc, "Table 5.1 Testing Summary", font="Times New Roman", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, spacing=1.0)
    add_table(
        doc,
        ["Test area", "Expected behavior", "Observed result"],
        [
            ["Sparse input validation", "At least two values required for prediction", "Backend rejects underspecified payloads and frontend warns the user"],
            ["Prediction result", "Risk, confidence, and label should be shown", "Result card renders risk percent, confidence, coverage, and care note"],
            ["History persistence", "Saved predictions should reload later", "Excel workbook entries are restored on dashboard refresh"],
            ["Diet personalization", "Latest profile should influence recommendations", "Diet page reads saved history and builds meal guidance"],
            ["Fallback resilience", "Frontend should not fail silently on backend outage", "Local estimate path remains available for some failures"],
        ],
        [Inches(1.6), Inches(2.3), Inches(2.5)],
    )
    add_paragraph(doc, "Overall, the testing results show that the current HeartGuard prototype is strongest as an explainable, continuity-oriented web workflow. Its validation strengths lie in interface stability, saved history reuse, and support-module linkage rather than in direct claims of hospital-grade deployment readiness.")
    doc.add_page_break()

    # Chapter 6
    add_chapter_title(doc, "6", "RESULTS AND DISCUSSION")
    add_section_title(doc, "6.1 Project Interface and Output")
    result_paragraphs = [
        "The most important implementation result is that HeartGuard produces an understandable prediction workflow that remains usable even when a patient cannot provide every health field. This is a meaningful improvement over rigid intake designs because preventive screening often begins with partial information rather than full laboratory records.",
        "Another important result is continuity after prediction. The backend saves successful outputs into workbook history, and the dashboard reconstructs that history after refresh instead of leaving the user with an empty screen. This creates a more trustworthy experience and makes the forecast useful beyond a single session.",
        "The surrounding modules add practical value. The diet page transforms the latest prediction into meal strategies, the chatbot provides accessible health guidance, the monitoring page visualizes vitals, the mood page records emotional well-being, and the medications module helps structure adherence habits. Together, these modules extend the system from prediction into preventive support."
    ]
    for p in result_paragraphs:
        add_paragraph(doc, p)

    add_section_title(doc, "6.2 Prediction Continuity and Support Features")
    add_paragraph(doc, "Table 6.1 Results and Discussion Summary", font="Times New Roman", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, spacing=1.0)
    add_table(
        doc,
        ["Result area", "Observed value", "Discussion"],
        [
            ["Prediction workflow", "Sparse-input scoring works with confidence feedback", "Suitable for demonstration and early preventive interaction"],
            ["Dashboard continuity", "Saved records reload from workbook history", "Improves trust and supports longitudinal review"],
            ["Support modules", "Diet, monitoring, mood, meds, chat, and community are integrated around prediction", "Transforms risk insight into broader preventive action"],
            ["Research assets", "Notebook metrics and pkl model files are preserved in the repo", "Supports future transition toward stronger model-serving integration"],
        ],
        [Inches(1.6), Inches(1.9), Inches(2.9)],
    )

    add_section_title(doc, "6.3 Implementation Tradeoffs")
    tradeoff_paragraphs = [
        "The present design reflects an intentional tradeoff between predictive sophistication and explainable usability. Although the repository contains model artifacts and notebook experimentation, the live public prototype does not currently expose a direct black-box classifier. Instead, it uses weighted clinical scoring so that the system remains understandable and stable under partial inputs.",
        "A second tradeoff concerns persistence. Excel-backed storage is easy to audit and sufficient for classroom or single-team use, but it is not durable enough for large-scale or multi-user deployment. This is clearly acknowledged in the deployment notes and should be addressed before production use.",
        "These tradeoffs do not reduce the academic value of the project. Instead, they help position HeartGuard accurately: it is a credible and well-connected prototype for preventive, readmission-aware support, while still requiring database hardening, model-serving alignment, and broader validation for future healthcare deployment."
    ]
    for p in tradeoff_paragraphs:
        add_paragraph(doc, p)
    doc.add_page_break()

    # Chapter 7
    add_chapter_title(doc, "7", "CONCLUSION")
    conclusion = (
        "HeartGuard demonstrates that heart-failure-aware forecasting can be made more useful when prediction is combined with continuity, explainability, and supportive follow-up modules. By integrating multi-source healthcare datasets, preserving notebook-based model experimentation, and delivering a deployed web prototype with risk scoring, prediction history, dashboard summaries, diet guidance, chatbot assistance, and wellness support features, the project moves beyond a narrow machine-learning exercise.\n\n"
        "The report also shows that the current implementation is honest about its scope. It is an academic system designed for preventive insight generation, not yet a hospital-certified readmission service. Even so, it contributes meaningful value by showing how forecasting, persistence, and user support can be connected in a single coherent workflow."
    )
    for block in conclusion.split("\n\n"):
        add_paragraph(doc, block)
    doc.add_page_break()

    # Chapter 8
    add_chapter_title(doc, "8", "FUTURE WORK")
    future_work = [
        "Integrate the deployed web interface with validated model-serving pipelines derived from the stored notebook and pkl artifacts.",
        "Replace Excel persistence with SQLite or PostgreSQL for durable multi-user history and concurrency support.",
        "Add secure authentication and role-based dashboards for patients, clinicians, and administrators.",
        "Include prospective clinical validation, calibration studies, and fairness analysis across diverse patient groups.",
        "Extend wellness support with reminders, notifications, and device-linked monitoring inputs.",
    ]
    add_bullets(doc, future_work)

    add_paragraph(doc, "REFERENCES", font="Times New Roman", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=14, after=10, spacing=1.0)
    references = [
        '[1] H. Hajishah et al., "Evaluation of machine learning methods for prediction of heart failure mortality and readmission: meta-analysis," BMC Cardiovascular Disorders, vol. 25, 2025.',
        '[2] J.-S. Rachoin, K. Hunter, J. Varallo, and E. Cerceo, "Impact of time from discharge to readmission on outcomes," BMJ Open, vol. 14, 2024.',
        '[3] R. O. Baris and C. E. Tabit, "Heart Failure Readmission Prevention Strategies," Journal of Clinical Medicine, vol. 14, no. 16, 2025.',
        '[4] R. B. D\'Agostino Sr. et al., "General cardiovascular risk profile for use in primary care: the Framingham Heart Study," Circulation, vol. 117, no. 6, 2008.',
        '[5] Project repository artifacts: FastAPI backend, workbook persistence, processed datasets, and notebook-based model training assets in the HeartGuard workspace.',
    ]
    for ref in references:
        add_paragraph(doc, ref, size=11, after=3, spacing=1.0)
    doc.add_page_break()

    add_paragraph(doc, "ANNEXURE", font="Times New Roman", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10, spacing=1.0)
    add_paragraph(doc, "Table A.1 Dataset and Variable Summary", font="Times New Roman", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=4, spacing=1.0)
    add_table(
        doc,
        ["Dataset / Variable group", "Current project evidence", "Use in HeartGuard"],
        [
            [f"Master dataset ({dataset_counts['master_rows']:,} x {dataset_counts['master_cols']})", "Combined harmonized records", "Broad research foundation"],
            [f"Cardiovascular dataset ({dataset_counts['cardio_rows']:,} x {dataset_counts['cardio_cols']})", "General heart-risk features", "Preventive risk context"],
            [f"Diabetes dataset ({dataset_counts['diabetes_rows']:,} x {dataset_counts['diabetes_cols']})", "Metabolic comorbidity signals", "Diet and chronic-risk context"],
            [f"Framingham dataset ({dataset_counts['framingham_rows']:,} x {dataset_counts['framingham_cols']})", "Classic long-term cardiovascular variables", "Primary-care risk framing"],
            [f"Heart failure dataset ({dataset_counts['heart_failure_rows']:,} x {dataset_counts['heart_failure_cols']})", "Advanced variables like EF and creatinine", "Readmission-aware sensitivity"],
            ["Deployed predictor fields", "age, sex, BMI, cholesterol, BP, glucose, activity, smoking, hypertension, diabetes, heart rate, EF, creatinine", "Explainable sparse-input scoring"],
        ],
        [Inches(2.25), Inches(2.15), Inches(2.1)],
    )
    add_paragraph(doc, "The annexure highlights an important point about the current project: HeartGuard is not based on a single isolated dataset or a single monolithic model idea. Instead, it combines a practical web workflow with a broader research base that can support future improvements in predictive fidelity, personalization, and deployment maturity.")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved report to {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
